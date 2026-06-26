# 수정된 소개서 docx 파일의 내용을 UTF-8 텍스트로 추출
from docx import Document
import os

doc = Document(r'd:\antigravity_vibecoding\BT 3바디 ai테스트\CODEMAP_AI_소개서_수정.docx')

output = []
output.append("=== PARAGRAPHS ===\n")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        output.append(f"[{i}] {para.text}")

for j, table in enumerate(doc.tables):
    output.append(f"\n=== TABLE {j} ===")
    for row in table.rows:
        cells = [cell.text.replace('\n', ' | ') for cell in row.cells]
        output.append(' || '.join(cells))

out_path = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docx_extract.txt'
with open(out_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))

print('Done:', out_path)
