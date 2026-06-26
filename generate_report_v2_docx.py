# health_report_v2.md를 프리미엄 스타일의 워드 문서(.docx)로 변환하는 스크립트
import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = docx.Document()

# 기본 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# 페이지 여백 설정
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 색상 정의
PRIMARY = RGBColor(79, 70, 229)
SECONDARY = RGBColor(59, 130, 246)
DARK = RGBColor(30, 41, 59)
MUTED = RGBColor(100, 116, 139)
ACCENT = RGBColor(245, 158, 11)
DANGER = RGBColor(220, 38, 38)
SUCCESS = RGBColor(16, 185, 129)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4F46E5')
    
    # 데이터
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            # 짝수 행 배경
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F8FAFC')
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table

def add_insight_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('📊 통계 분석: ')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run = p.add_run(text)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED
    run.font.italic = True

# =====================================================================
# 표지
# =====================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CODEMAP AI 대국민 3Body 통합 건강 통계 보고서')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(14)
run.font.color.rgb = SECONDARY
run.font.bold = True

doc.add_paragraph()

title = doc.add_heading('대한민국 건강 보고서', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Malgun Gothic'
    run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('제2판 (V2)')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = ACCENT

doc.add_paragraph()

info_lines = [
    '분석 대상: 고유 사용자 3,267명 (V1 대비 +1,312명, 67.1% 증가)',
    '기준 일자: 2026년 6월 17일',
    '데이터: members_v4 컬렉션 (AI분석 완료 기준)',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

doc.add_page_break()

# =====================================================================
# 1. 정제 데이터 개요
# =====================================================================
h1 = doc.add_heading('1. 정제 데이터 개요 및 이용 현황', level=1)
for run in h1.runs:
    run.font.name = 'Malgun Gothic'

p = doc.add_paragraph('서버의 라이브 데이터를 훼손하지 않는 안전한 PC 로컬 가공 처리를 통해 총 6,272건의 레코드를 정제했습니다.')

bullets = [
    ('전체 데이터', '6,272건'),
    ('분석 미완료(Pending) 데이터 제외', '2,820건'),
    ('테스트/교육용 데이터 제외', '32건'),
    ('동일인 중복 제거', '153건'),
    ('최종 정제 완료 고유 회원', '3,267명'),
]
for label, val in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.font.bold = True
    run.font.name = 'Malgun Gothic'
    p.add_run(val).font.name = 'Malgun Gothic'

doc.add_paragraph()

add_styled_table(doc,
    ['이용 환경 구분', '고유 사용자 수', '비율'],
    [
        ['PC 버전 (지점 설치형)', '1,148명', '35.1%'],
        ['LITE 버전 (행사·야외·온라인용)', '1,761명', '53.9%'],
        ['기타', '358명', '11.0%'],
        ['합계', '3,267명', '100.0%'],
    ],
    [6, 4, 3]
)

add_insight_box(doc, 'V1 대비 LITE 버전 유입 비중이 42.0%에서 53.9%로 상승했습니다. 전국 건강 캠페인을 통한 일반 국민 참여가 급증하고 있습니다.')

# =====================================================================
# 2. 인구통계
# =====================================================================
doc.add_page_break()
h2 = doc.add_heading('2. 체험 회원 인구통계학적 분포', level=1)
for run in h2.runs:
    run.font.name = 'Malgun Gothic'

add_styled_table(doc,
    ['연령대', '남성 (명)', '여성 (명)', '합계 (명)', '비율'],
    [
        ['10대 이하', '—', '—', '31', '0.9%'],
        ['20대', '—', '—', '109', '3.3%'],
        ['30대', '—', '—', '202', '6.2%'],
        ['40대', '—', '—', '341', '10.4%'],
        ['50대', '—', '—', '790', '24.2%'],
        ['60대', '—', '—', '1,006', '30.8%'],
        ['70대', '—', '—', '640', '19.6%'],
        ['80대 이상', '—', '—', '148', '4.5%'],
        ['합계', '826', '2,441', '3,267', '100.0%'],
    ],
    [3, 2.5, 2.5, 2.5, 2.5]
)

add_insight_box(doc, '여성 비율 74.7%(2,441명), 50대(24.2%)와 60대(30.8%)가 전체의 55.0%를 차지합니다. 70대 이상 고령층 참여가 V1 대비 확대되었습니다.')

# =====================================================================
# 3. 3Body 분석
# =====================================================================
doc.add_page_break()
h3 = doc.add_heading('3. 3Body (몸·마음·뇌) 통합 건강 지표 심층 분석', level=1)
for run in h3.runs:
    run.font.name = 'Malgun Gothic'

doc.add_heading('A. 성별 3Body 평균 점수 비교', level=2)

add_styled_table(doc,
    ['성별', '신체 (Body)', '뇌 (Brain)', '마음 (Mind)'],
    [
        ['남성 (826명)', '71.1점', '77.2점', '59.5점'],
        ['여성 (2,441명)', '70.3점', '77.3점', '58.3점'],
        ['전체 평균 (3,267명)', '70.5점', '77.3점', '58.6점'],
    ],
    [5, 3, 3, 3]
)

add_insight_box(doc, '성별에 관계없이 마음(Mind) 점수가 58~59점대로 가장 낮고, 뇌(Brain) 점수가 77점대로 상대적으로 높게 분포됩니다.')

doc.add_paragraph()
doc.add_heading('B. 연령대별 3Body 평균 점수 비교', level=2)

add_styled_table(doc,
    ['연령대', '인원', '신체 (Body)', '뇌 (Brain)', '마음 (Mind)'],
    [
        ['10대 이하', '31명', '73.1점', '63.3점', '54.7점'],
        ['20대', '109명', '76.3점', '71.9점', '59.9점'],
        ['30대', '202명', '75.4점', '76.3점', '58.0점'],
        ['40대', '341명', '75.4점', '80.0점', '60.5점'],
        ['50대', '790명', '74.1점', '79.4점', '60.2점'],
        ['60대', '1,006명', '69.3점', '76.9점', '58.4점'],
        ['70대', '640명', '64.5점', '76.9점', '56.8점'],
        ['80대 이상', '148명', '61.7점', '74.4점', '55.4점'],
    ],
    [3, 2, 2.5, 2.5, 2.5]
)

add_insight_box(doc, '신체는 20대(76.3)에서 정점 후 완만 감소. 뇌는 40대(80.0)에서 정점. 마음은 전 연령대 60점 미만으로 만성 저조하여 최대 과제입니다.')

# =====================================================================
# 4. 7Code 분석
# =====================================================================
doc.add_page_break()
h4 = doc.add_heading('4. 7Code 웰니스 에너지 지표 상세 분석', level=1)
for run in h4.runs:
    run.font.name = 'Malgun Gothic'

doc.add_heading('A. 성별 7Code 평균 점수', level=2)

add_styled_table(doc,
    ['구분', '1코드', '2코드', '3코드', '4코드', '5코드', '6코드', '7코드'],
    [
        ['남성', '54.6', '62.3', '61.7', '66.7', '63.5', '63.8', '69.3'],
        ['여성', '55.3', '60.7', '60.8', '64.4', '62.8', '62.3', '68.4'],
    ],
    [2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
)

doc.add_paragraph()
doc.add_heading('B. 7Code 취약 코드 분석 (개인별 최저 점수 코드)', level=2)

add_styled_table(doc,
    ['순위', '코드 (증상)', '해당 인원', '비율'],
    [
        ['1위', '1코드 (힐링 라이프 - 이완 장애 및 휴식 결핍)', '1,343명', '41.1%'],
        ['2위', '5코드 (집중 & 직관 - 인지 피로 및 집중 저하)', '491명', '15.0%'],
        ['3위', '6코드 (통찰 & 지혜 - 스트레스성 두뇌 기능 저하)', '416명', '12.7%'],
        ['4위', '2코드 (스트레스 관리 - 정서적 누적 긴장)', '373명', '11.4%'],
        ['5위', '3코드 (감정 균형 - 감정 기복 및 무기력)', '286명', '8.8%'],
        ['6위', '4코드 (활력 충전 - 육체 에너지 고갈)', '210명', '6.4%'],
        ['7위', '7코드 (의식 & 영성 - 자아 정체성 혼란)', '137명', '4.2%'],
    ],
    [1.5, 7, 2, 2]
)

add_insight_box(doc, '국민 10명 중 4명(41.1%)이 1코드(이완 장애) 최취약. V1(38.4%) 대비 2.7%p 상승하여, 이완 능력 결핍이 가장 보편적 취약점임이 재확인되었습니다.')

# =====================================================================
# 5. 자세 분석
# =====================================================================
doc.add_page_break()
h5 = doc.add_heading('5. 신체 관절 정렬 및 자세 대칭성 분석 (N=3,008명)', level=1)
for run in h5.runs:
    run.font.name = 'Malgun Gothic'

add_styled_table(doc,
    ['자세 측정 지표', 'Good (정상)', 'Fair (주의)', 'Poor (위험)', '표본'],
    [
        ['거북목 (FHP) 및 경추 정렬', '0.5%', '79.1%', '20.4%', '3,008명'],
        ['어깨 / 골반 좌우 대칭', '59.4%', '38.7%', '1.9%', '3,008명'],
        ['측면 척추 정렬 (흉추/요추)', '3.4%', '77.0%', '19.6%', '3,008명'],
        ['하체 기저면 (무릎/다리/발목)', '90.0%', '8.0%', '2.0%', '3,008명'],
        ['귀-어깨-고관절-무릎 수직선 이탈', '1.4%', '76.2%', '22.4%', '3,008명'],
    ],
    [5, 2, 2, 2, 2]
)

add_insight_box(doc, 'V1 대비 거북목 Poor 15.4%→20.4%(+5.0%p), 수직선 이탈 Poor 15.3%→22.4%(+7.1%p). 국민 5명 중 1명 이상이 전신 수직 정렬 위험 상태입니다.')

# =====================================================================
# 6. 운동 능력
# =====================================================================
doc.add_paragraph()
h6 = doc.add_heading('6. 연령대별 운동 능력 비교', level=1)
for run in h6.runs:
    run.font.name = 'Malgun Gothic'

add_styled_table(doc,
    ['연령대', '스쿼트 (회/점수)', '푸쉬업 (회/점수)', '유연성', '팔올리기', '인원'],
    [
        ['10대이하', '8.6회 (74점)', '6.8회 (66점)', '72.1점', '69.6점', '31명'],
        ['20대', '7.3회 (75점)', '7.5회 (66점)', '79.9점', '72.6점', '109명'],
        ['30대', '8.1회 (74점)', '7.5회 (63점)', '76.2점', '71.8점', '202명'],
        ['40대', '8.3회 (73점)', '8.4회 (65점)', '79.4점', '73.6점', '341명'],
        ['50대', '8.4회 (71점)', '8.2회 (63점)', '78.1점', '71.4점', '790명'],
        ['60대', '6.9회 (67점)', '6.3회 (61점)', '78.2점', '69.6점', '1,006명'],
        ['70대', '6.6회 (68점)', '6.7회 (62점)', '77.9점', '67.3점', '640명'],
        ['80대이상', '5.2회 (65점)', '4.6회 (52점)', '77.3점', '65.5점', '148명'],
    ],
    [2.5, 2.5, 2.5, 2, 2, 1.5]
)

add_insight_box(doc, '하체 근력은 50대까지 8회 이상 유지 후 60대부터 급감. 유연성은 전 연령대 72~80점대로 균일. 80대 푸쉬업 자세 점수(52점)가 유일하게 60점 미만입니다.')

# =====================================================================
# 7. 건강나이
# =====================================================================
doc.add_page_break()
h7 = doc.add_heading('7. 건강나이(웰니스 에이지) 연령대별 추정 분석', level=1)
for run in h7.runs:
    run.font.name = 'Malgun Gothic'

p = doc.add_paragraph('AI가 추정한 신체나이, 뇌나이, 마음나이, 얼굴나이를 실제 나이와 비교한 결과입니다.')

add_styled_table(doc,
    ['연령대', '실제 나이', '신체나이', '뇌나이', '마음나이', '얼굴나이', '종합나이'],
    [
        ['20대', '25.6세', '26.9세', '35.8세', '35.9세', '25.0세', '30.7세'],
        ['30대', '34.0세', '33.4세', '37.4세', '47.4세', '32.3세', '36.4세'],
        ['40대', '45.1세', '43.4세', '41.4세', '56.7세', '42.1세', '44.5세'],
        ['50대', '55.0세', '53.5세', '47.3세', '66.6세', '51.9세', '53.3세'],
        ['60대', '64.6세', '65.1세', '54.8세', '75.0세', '61.8세', '63.0세'],
        ['70대', '73.8세', '74.5세', '58.7세', '81.3세', '71.0세', '70.2세'],
    ],
    [2.2, 2, 2, 2, 2, 2, 2]
)

add_insight_box(doc, '뇌나이는 실제 나이보다 10~15세 젊고, 마음나이는 항상 10세 이상 높습니다. 현대인의 뇌는 젊지만 마음은 늙어가고 있습니다.')

# =====================================================================
# 8. 에너지 순환
# =====================================================================
doc.add_paragraph()
h8 = doc.add_heading('8. 3Body 에너지 순환 상태 분석', level=1)
for run in h8.runs:
    run.font.name = 'Malgun Gothic'

bullets_energy = [
    ('에너지 순환 관리 필요도 [높음]', '3,179명 (97.3%)'),
    ('에너지 순환 관리 필요도 [매우 높음]', '68명 (2.1%)'),
    ('에너지 순환 관리 필요도 [보통]', '12명 (0.4%)'),
    ('에너지 순환 관리 필요도 [낮음]', '8명 (0.2%)'),
]
for label, val in bullets_energy:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.font.bold = True
    run.font.name = 'Malgun Gothic'
    p.add_run(val).font.name = 'Malgun Gothic'

add_insight_box(doc, '99.4%가 에너지 순환 집중 관리군에 해당합니다. 대한민국 국민 절대 다수가 자율신경계 과긴장과 에너지 순환 정체 상태에 놓여있습니다.')

# =====================================================================
# 9. V1 vs V2
# =====================================================================
doc.add_page_break()
h9 = doc.add_heading('9. V1 vs V2 핵심 변화 비교', level=1)
for run in h9.runs:
    run.font.name = 'Malgun Gothic'

add_styled_table(doc,
    ['지표', '제1판 (1,955명)', '제2판 (3,267명)', '변화'],
    [
        ['표본 규모', '1,955명', '3,267명', '+67.1% 증가'],
        ['신체(Body) 평균', '71.4점', '70.5점', '▼ 0.9점'],
        ['뇌(Brain) 평균', '76.5점', '77.3점', '▲ 0.8점'],
        ['마음(Mind) 평균', '59.5점', '58.6점', '▼ 0.9점'],
        ['1코드 취약 비율', '38.4%', '41.1%', '▲ 2.7%p'],
        ['거북목 Poor 비율', '15.4%', '20.4%', '▲ 5.0%p'],
        ['수직선 이탈 Poor', '15.3%', '22.4%', '▲ 7.1%p'],
        ['에너지 관리 필요', '99.7%', '99.4%', '동일 수준'],
    ],
    [4, 3, 3, 3]
)

add_insight_box(doc, 'V1의 모든 핵심 패턴이 67% 확대된 표본에서 강화·재확인되었습니다. 자세 정렬 위험 비율이 V1 대비 5~7%p 상승하여 실제 국민 건강 실태가 더 정확하게 반영되고 있습니다.')

# =====================================================================
# 10. 결론
# =====================================================================
doc.add_paragraph()
h10 = doc.add_heading('10. 결론: 홀리스틱 웰니스와 대국민 100만 캠페인', level=1)
for run in h10.runs:
    run.font.name = 'Malgun Gothic'

p = doc.add_paragraph()
run = p.add_run('핵심 발견 사항')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = PRIMARY

findings = [
    ('마음이 가장 아프다', '전 연령대에서 마음(Mind) 점수가 55~60점대로 가장 낮고, 마음나이가 실제 나이보다 항상 10세 이상 높습니다.'),
    ('국민 10명 중 4명은 이완 불능', '1코드(이완 장애)가 41.1%로 가장 보편적인 취약점이며, 자율신경계 만성 과긴장 상태를 반영합니다.'),
    ('5명 중 1명 이상이 거북목 위험', '거북목(20.4%)과 전신 수직 정렬 이탈(22.4%)이 V1 대비 악화되었습니다.'),
    ('뇌는 가장 젊다', '뇌나이가 실제 나이보다 10~15세 젊어 뇌 건강의 회복 잠재력이 가장 높은 영역입니다.'),
    ('99.4%가 에너지 순환 관리 필요', '일반 국민 절대 다수가 에너지 정체 상태에 있습니다.'),
]

for title, desc in findings:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title} — ')
    run.font.bold = True
    run.font.name = 'Malgun Gothic'
    run = p.add_run(desc)
    run.font.name = 'Malgun Gothic'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('건강은 단순히 질병이 없는 상태가 아니라, 몸의 정렬과 마음의 이완, 뇌의 활력이 서로 조화를 이루는 홀리스틱 웰니스(Holistic Wellness) 관점에서 다루어져야 합니다.')
run.font.name = 'Malgun Gothic'
run.font.italic = True
run.font.color.rgb = DARK

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CODEMAP AI 대국민 100만 캠페인')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('대한민국 국민 100만 명이 자신만의 CODEMAP을 지니고\n균형 있는 몸·마음·뇌를 관리하는 건강한 대한민국을 구현합니다.')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(11)
run.font.color.rgb = MUTED

doc.add_paragraph()
doc.add_paragraph()

# 법적 고지
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('본 보고서의 모든 데이터는 비의료 건강관리서비스 가이드라인을 준수하며, 의료행위를 대체하지 않습니다.')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(8)
run.font.color.rgb = MUTED
run.font.italic = True

# 저장
output = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\대한민국_건강보고서_V2.docx'
doc.save(output)
print(f'Word document saved: {output}')
