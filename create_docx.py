from docx import Document
import os

doc = Document()

# Title
heading = doc.add_heading('BTC 3바디 AI 분석기 V5.0 - 지점 설치 및 세팅 매뉴얼', 0)

doc.add_paragraph('안녕하세요! BTC 3바디 AI 분석기 V5.0 지점 원장님(관리자)용 설치 매뉴얼입니다.\n본 가이드는 지점 PC에 프로그램을 설치하고, 최적의 진단 환경을 세팅하는 과정을 안내합니다.')

# Update Highlight
p = doc.add_paragraph()
p.add_run('[V5.0 주요 업데이트 사항]\n').bold = True
p.add_run('- 마우스 기반 뇌 인지 테스트(TMT) 도입: 기존의 카메라 기반 반응속도 테스트가 정밀한 마우스 클릭 방식으로 전면 교체되었습니다. 조명이나 거리 등 환경적 제약 없이 정확한 \'뇌 나이\' 평가가 가능합니다.\n')
p.add_run('- 설치형 프로그램(Installer): 무설치 압축 파일 방식에서 윈도우 자동 설치 마법사 방식으로 변경되어 안정성과 바탕화면 바로가기 접근성이 강화되었습니다.\n')
p.add_run('- 경량화 및 속도 향상: 구글 로그인 연동이 삭제되고, 내부 로직이 최적화되어 더욱 빠르고 쾌적하게 동작합니다.')

doc.add_heading('1. 권장 PC 사양 및 측정 환경', level=1)
doc.add_paragraph('정확한 AI 진단(카메라 비전 추적 및 인지 테스트)을 위해 다음 환경을 권장합니다.')
doc.add_paragraph('권장 사양: 인텔 코어 i3 (8세대 이상), 메모리 8GB 이상', style='List Bullet')
doc.add_paragraph('입력 장치: 마우스 필수 (새로운 TMT 인지 반응속도 테스트에 사용됩니다.)', style='List Bullet')
doc.add_paragraph('카메라(웹캠) 및 공간:\n  - 신체 밸런스 및 관상 분석에는 여전히 카메라가 사용됩니다.\n  - 설치 높이: 바닥면에서 약 1m ~ 1.2m (성인 배꼽 높이)에 수평으로 고정\n  - 측정 거리: 거리에 따라 광각 웹캠(FOV 110도 이상, 예: 앱코 APC930) 사용을 강력히 권장합니다.', style='List Bullet')
doc.add_paragraph('🚨 주의사항: 진단 중 엑셀, 카카오톡, 인터넷 창 등은 메모리 부족을 유발할 수 있으므로 가급적 닫아주세요.', style='List Bullet')

doc.add_heading('2. 프로그램 설치 방법', level=1)
doc.add_paragraph('1. 본사에서 제공한 BTC_3Body_AI_Analyzer_v5.0.0.exe 설치 파일을 다운로드하여 더블클릭합니다.')
doc.add_paragraph('2. 설치 마법사가 실행되면 지시에 따라 [다음]을 눌러 설치를 완료합니다.')
doc.add_paragraph('3. 설치가 완료되면 바탕화면에 "BT 3바디 AI 진단" 바로가기 아이콘이 생성됩니다.')

p2 = doc.add_paragraph()
p2.add_run('[스마트 앱 컨트롤 / Windows SmartScreen 차단 시 해결법]\n').bold = True
p2.add_run('사내 배포용 앱이므로 윈도우 보안 경고가 뜰 수 있습니다.\n')
p2.add_run('- 파란색 경고창: 얇은 글씨인 [추가 정보] 클릭 후 [실행] 버튼 클릭\n')
p2.add_run('- 회색 팝업 차단 시: 윈도우 검색창에 \'스마트 앱 컨트롤\' 검색 후 설정을 [끄기]로 변경')

doc.add_heading('3. 최초 1회 시스템 인증 (하드웨어 락)', level=1)
doc.add_paragraph('본 프로그램은 고가의 AI 라이선스 보호를 위해 \'1 PC = 1 라이선스\' 정책을 사용합니다.\n(최초 1회만 인증하면 이후에는 자동으로 실행됩니다.)')
doc.add_paragraph('1. 바탕화면의 아이콘을 더블클릭하여 프로그램을 실행합니다.')
doc.add_paragraph('2. "지점 인증 및 등록" 화면이 나타나면, 아래 정보를 정확히 입력합니다:\n   - 지점명 / 관리자명 / 연락처 / 설치 담당자')
doc.add_paragraph('3. 지점 인증번호: 본사 마스터 암호인 BTC15771785 를 입력합니다.')
doc.add_paragraph('4. [인증하기] 버튼을 누르면 해당 PC의 메인보드 고유번호가 중앙 서버에 등록되며 라이선스가 영구 귀속됩니다.')

p3 = doc.add_paragraph()
p3.add_run('[무단 복제 및 이동 설치 절대 금지]\n').bold = True
p3.add_run('인증이 완료된 폴더를 USB 등으로 복사해 다른 PC에서 실행할 경우, 해킹으로 간주되어 프로그램이 영구 차단됩니다. PC 교체 시에는 반드시 본사에 기존 기기 라이선스 해지를 요청하셔야 합니다.')

doc.add_heading('4. 진단 시작 및 진행 가이드', level=1)
doc.add_paragraph('인증이 완료되면 "음성 안내 활성화" 확인 후 [진단 센터 입장하기]를 눌러 시작합니다.')
doc.add_paragraph('1. 회원 등록: [신규 진단 시작하기]를 누르고 회원의 이름, 나이, 성별을 간략히 입력합니다.')
doc.add_paragraph('2. 신체 밸런스 테스트 (카메라 사용):\n   - 정면/측면 균형, 스쿼트, 한발 서기 등 화면의 타이머와 가이드에 맞춰 자세를 취합니다.')
doc.add_paragraph('3. [NEW] 뇌 인지 반응속도 테스트 (마우스 사용):\n   - 화면에 나타나는 3색 타겟(공)을 순서대로 빠르게 마우스로 클릭합니다. (총 2라운드, 15초 제한)\n   - 1라운드: 오름차순 (1→10)\n   - 2라운드: 내림차순 (10→1)')
doc.add_paragraph('4. 안면 노화 분석: 마지막 단계에서 카메라 앞으로 다가와 얼굴을 인식시킵니다.')
doc.add_paragraph('5. 결과 리포트: 신체 연령, 뇌 나이, 3바디/7차크라 분석 결과 및 맞춤 솔루션을 확인합니다.')

p4 = doc.add_paragraph()
p4.add_run('\n지점에서 BT 3바디 AI 분석기 V5.0을 통해 한 차원 높은 전문 상담을 제공해 보시기 바랍니다!').bold = True

file_path = os.path.join(os.getcwd(), 'BTC_3바디_AI분석기_v5.0_지점설치매뉴얼.docx')
doc.save(file_path)
print(f'Saved to {file_path}')
