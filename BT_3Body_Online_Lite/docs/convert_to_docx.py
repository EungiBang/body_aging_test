from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ===== 페이지 설정 =====
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== 스타일 설정 =====
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

h1 = doc.styles['Heading 1']
h1.font.name = '맑은 고딕'
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

h2 = doc.styles['Heading 2']
h2.font.name = '맑은 고딕'
h2.font.size = Pt(15)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2d, 0x3a, 0x8c)
h2.paragraph_format.space_before = Pt(20)
h2.paragraph_format.space_after = Pt(10)

h3 = doc.styles['Heading 3']
h3.font.name = '맑은 고딕'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x4a, 0x4a, 0x8a)
h3.paragraph_format.space_before = Pt(14)
h3.paragraph_format.space_after = Pt(6)

# ===== 유틸리티 함수 =====
def add_rich(doc, text, indent=0, bullet=''):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    full = (bullet + ' ' if bullet else '') + text
    parts = re.split(r'(\*\*.*?\*\*)', full)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(9.5)
                run.bold = True
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(9)
    doc.add_paragraph('')
    return table

# ===== 표지 =====
for _ in range(5):
    doc.add_paragraph('')

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('BTC 3바디 AI 분석기 V4.1.0')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = '맑은 고딕'

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('사용 및 평가기준 매뉴얼')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x8a)
run.font.name = '맑은 고딕'

doc.add_paragraph('')

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc_p.add_run('11단계 전체 웰니스 스크리닝 프로세스\n단계별 채점 기준 및 종합 나이 산출 공식')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '맑은 고딕'

doc.add_paragraph('')
doc.add_paragraph('')

org_p = doc.add_paragraph()
org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = org_p.add_run('브레인트레이닝센터(BTC)')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2d, 0x3a, 0x8c)
run.font.name = '맑은 고딕'

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf_p.add_run('[ 대외비 - 지점 관리자 전용 ]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
run.font.name = '맑은 고딕'
run.font.bold = True

doc.add_page_break()

# ===== 서문 =====
add_rich(doc, '이 매뉴얼은 V4.1.0에 탑재된 **11단계 전체 웰니스 스크리닝 프로세스**를 지점 관리자가 완벽히 숙지하고, 회원들에게 정확한 분석 평가 기준을 설명할 수 있도록 제작되었습니다.')

# ===== Chapter 1 =====
doc.add_heading('Chapter 1. 완벽한 분석을 위한 올바른 가이드라인', level=2)
add_rich(doc, 'AI가 가장 정확한 데이터를 추출할 수 있도록 사전 세팅과 회원 안내에 신경 써주세요.')

doc.add_heading('1-1. 기기 및 환경 세팅', level=3)
add_rich(doc, '**카메라 위치:** 사용자의 눈높이에 맞추어 설치하되, 사용자가 카메라 앞 **1.5m ~ 2m 거리**에 섰을 때 **머리끝부터 발끝까지 전신이 모두 화면에 들어와야** 합니다.', bullet='*')
add_rich(doc, '**조명:** 역광을 피하고, 피사체(회원)의 얼굴과 몸이 밝게 비춰지는 순행광 위치가 가장 좋습니다.', bullet='*')
add_rich(doc, '**PC 환경:** AI 엔진이 램(RAM)을 집중적으로 사용하므로, 측정 전 **카카오톡, 무거운 엑셀 창, 불필요한 크롬 탭은 반드시 닫아주세요.**', bullet='*')

doc.add_heading('1-2. 회원의 복장 및 준비 자세', level=3)
add_rich(doc, '**복장:** 헐렁한 옷이나 어두운 계열의 옷은 관절 포인트를 AI가 오인할 수 있습니다. 체형이 잘 드러나는 **밝은 톤의 수련복/운동복** 착용을 적극 권장합니다.', bullet='*')
add_rich(doc, '**자세 안내:** "자연스럽게 서주세요" 보다는 **"평소 편하게 서있는 자세 그대로 서보세요"**라고 안내해야 합니다.', bullet='*')

# ===== Chapter 2 =====
doc.add_heading('Chapter 2. 단계별 분석 방법 및 점수 평가 기준', level=2)
add_rich(doc, '회원이 "제 점수가 어떻게 깎인 건가요?"라고 질문할 때, 아래의 기준을 바탕으로 명확히 설명해 주시면 신뢰도를 크게 높일 수 있습니다.')

# --- 1단계 ---
doc.add_heading('[1단계] 정면 자세 분석', level=3)
add_rich(doc, 'AI가 정면 사진에서 관절 좌표를 추출하여 **상체-하체의 좌우 대칭**을 수치로 계산합니다.')
add_rich(doc, '**어깨 기울기:** 0도 = 완벽 대칭 / 3도 이상 = 비대칭 주의', bullet='*')
add_rich(doc, '**골반 기울기:** 0도 = 완벽 대칭 / 3도 이상 = 비대칭 주의', bullet='*')
add_rich(doc, '**어깨-골반 너비 비율:** 1.2~1.4 = 표준 / 0.9 이하 = 비만 체형 경향 / 1.5 이상 = 역삼각형', bullet='*')
add_rich(doc, '**다리 형태:** 정상 / O자 / X자 -- 편차 각도 기반 판정', bullet='*')
add_rich(doc, '**무릎 정렬:** 정상 / 비대칭', bullet='*')

# --- 2단계 ---
doc.add_heading('[2단계] 측면 자세 분석', level=3)
add_rich(doc, '측면 사진에서 **거북목, 등 굽음, 라운드 숄더** 등을 수치로 계산합니다.')
add_rich(doc, '**거북목(FHP) 기울기:** 0~15도 정상 / 15도 이상 = 거북목', bullet='*')
add_rich(doc, '**상체(흉추) 기울기:** 0~10도 정상 / 10도 이상 = 기울어짐', bullet='*')
add_rich(doc, '**라운드 숄더 각도:** 0~5도 정상 / 10도 이상 = 어깨 말림', bullet='*')
add_rich(doc, '**등 굽힘(흉추 후만):** 0~10도 정상 / 15도 이상 = 주의 / 25도 이상 = 심한 굽음', bullet='*')
add_rich(doc, '**감점 원리:** 자세가 바르지 않으면 기본 60점대(+5살), 매우 우수하면 90점 이상(-10살)', bullet='*')

# --- 3단계 ---
doc.add_heading('[3단계] 균형 테스트 (눈 감고 한발 서기 15초)', level=3)
add_rich(doc, '눈을 감은 상태에서 한 발로 15초간 서는 테스트입니다. **전정기관(평형감각)과 고유수용감각**의 통합 능력을 평가합니다.')
add_rich(doc, '**발 내려놓은 횟수 (footDrops):** 15초 동안 든 발을 바닥에 내려놓은 횟수', indent=0.5)
add_rich(doc, '**흔들림 점수 (swayScore):** AI 모션 트래킹으로 측정한 신체 흔들림 정도', indent=0.5)

add_rich(doc, '**점수 기준표 (100점 만점):**')
add_table(doc,
    ['footDrops', '흔들림 낮음\n(sway<20)', '흔들림 보통\n(sway<60)', '흔들림 높음\n(sway>=60)'],
    [
        ['0회', '100점', '85점', '75점'],
        ['1회', '70점', '60점', '50점'],
        ['2회', '40점', '40점', '40점'],
        ['3회 이상', '20점', '20점', '20점'],
    ])
add_rich(doc, '**눈을 뜨고 수행한 경우:** footDrops에 **+2회 패널티** 자동 가산', bullet='*')

# --- 4단계 ---
doc.add_heading('[4단계] 팔 올리기 (견관절 가동범위)', level=3)
add_rich(doc, '양팔을 최대한 높이 올려 **어깨 관절의 유연성**을 측정합니다.')
add_rich(doc, '**평균 올림 각도:** 180도가 완벽한 수직 (높을수록 좋음)', bullet='*')
add_rich(doc, '**팔과 귀 밀착도:** 팔을 올렸을 때 귀에 얼마나 가까운지', bullet='*')
add_rich(doc, '**팔꿈치 펴짐:** 팔꿈치가 굽어지면 감점', bullet='*')

# --- 5단계 ---
doc.add_heading('[5단계] 유연성 테스트 (전굴)', level=3)
add_rich(doc, '선 자세에서 상체를 앞으로 굽혀 **후면 사슬 근육의 유연성**을 측정합니다.')
add_rich(doc, '**손끝 닿는 위치:** 발목 위 / 발목 / 바닥 / 손바닥 (낮을수록 유연)', bullet='*')
add_rich(doc, '**무릎 펴짐:** 무릎이 굽어지면 감점 (반칙)', bullet='*')
add_rich(doc, '**허리 굽힘 여부:** 허리가 충분히 굽혀지지 않으면 감점', bullet='*')

# --- 6단계 ---
doc.add_heading('[6단계] 근력 테스트 - 스쿼트 (15초)', level=3)
add_rich(doc, '15초 동안 최대한 많은 스쿼트를 수행하여 **하체 순간 근력**을 측정합니다.')
add_rich(doc, '**나이대별 평균 횟수 기준 (15초):**')
add_table(doc,
    ['나이대', '남성 평균', '여성 평균'],
    [
        ['20대', '18회', '15회'],
        ['30대', '15회', '12회'],
        ['40대', '12회', '10회'],
        ['50대', '10회', '8회'],
        ['60대', '8회', '6회'],
        ['70대+', '6회', '4회'],
    ])
add_rich(doc, '**점수 산출:** 나이대별 평균 = 70점 기준. 평균의 1.5배 이상 = 100점. **1회 차이 = 1.8세 변화**', bullet='*')
add_rich(doc, '**자세 보정 (formScore):** 60점 이상 = 전체 인정 / 45~59점 = 80% 인정 / 44점 이하 = 60% 인정', bullet='*')
add_rich(doc, '**신체 나이 환산:** 90점+ = 실제나이-10세 / 80점+ = -5세 / 70점 = 동일 / 60점 = +5세 / 50점 = +10세 / 50점 미만 = +15세', bullet='*')

# --- 7단계 ---
doc.add_heading('[7단계] 근력 테스트 - 팔굽혀펴기 (15초)', level=3)
add_rich(doc, '15초 동안 최대한 많은 팔굽혀펴기를 수행하여 **상체 순간 근력**을 측정합니다.')
add_rich(doc, '**나이대별 평균 횟수 기준 (15초):**')
add_table(doc,
    ['나이대', '남성 평균', '여성 평균'],
    [
        ['20대', '18회', '13회'],
        ['30대', '15회', '10회'],
        ['40대', '12회', '8회'],
        ['50대', '9회', '6회'],
        ['60대', '7회', '4회'],
        ['70대+', '4회', '2회'],
    ])
add_rich(doc, '**무릎 보조 (변형 푸시업) 규정:**')
add_rich(doc, '70세 이상 여성: 무릎 대고 수행 = **정식 기준 적용 (패널티 없음)**', indent=0.8, bullet='*')
add_rich(doc, '60~69세 여성: 무릎 대고 수행 = **60% 인정 (부분 패널티)**', indent=0.8, bullet='*')
add_rich(doc, '그 외: 무릎 보조 시 **0회 처리** (정자세 1회보다 항상 불리)', indent=0.8, bullet='*')

# --- 8단계 ---
doc.add_heading('[8단계] 두뇌 인지 반응 테스트 (100점 만점)', level=3)
add_rich(doc, '화면에 나오는 지시에 따라 손을 드는 테스트입니다. (스트룹 효과 원리 적용)')
add_rich(doc, '**반응 속도 점수 (30점 만점):**')
add_table(doc,
    ['반응 시간', '점수', '판정'],
    [
        ['600ms 이하', '30점', '상위권 전두엽 활성도'],
        ['601~900ms', '20점', '보통'],
        ['901~1200ms', '10점', '저하'],
        ['1200ms 초과', '0점', '심각한 저하'],
    ])
add_rich(doc, '**오답 억제력 점수 (70점 만점):**')
add_table(doc,
    ['오답 수', '점수'],
    [
        ['0개', '70점 (만점)'],
        ['1개', '50점'],
        ['2개', '30점'],
        ['3개 이상', '10점'],
    ])

# --- 9단계 ---
doc.add_heading('[9단계] 단기 작업기억력 - 마트 장보기 (100점 만점)', level=3)
add_rich(doc, '3초 동안 화면에 나오는 물건과 가격을 기억하고 순서대로 터치하는 테스트입니다.')
add_rich(doc, '**기억력 스팬 점수 (50점 만점):**')
add_table(doc,
    ['정답 수', '점수'],
    [
        ['6개 이상', '50점'],
        ['5개', '40점'],
        ['4개', '30점'],
        ['3개', '20점'],
        ['2개', '10점'],
        ['1개 이하', '5점'],
    ])
add_rich(doc, '**가격 계산 점수 (30점):** 총금액 정답 = 30점, 오답 = 10점', bullet='*')
add_rich(doc, '**사칙연산 트릭 점수 (20점):** 2문제 정답 = 20점 / 1문제 = 10점 / 0문제 = 0점', bullet='*')

# --- 10단계 ---
doc.add_heading('[10단계] 얼굴 활력 나이 (안면 분석)', level=3)
add_rich(doc, '카메라를 통해 안면을 촬영하고, AI가 **안면 밝기(Luma)와 표정**을 분석합니다.')
add_rich(doc, '**안면 밝기 (Luma):** 피부의 전반적인 밝기와 광채', bullet='*')
add_rich(doc, '**표정 분석:** 미세 표정 근육(Blendshapes)의 긴장도', bullet='*')
add_rich(doc, '(※ 부가 서비스인 K-관상 기질 분석과는 별개의 건강 활력 지표입니다.)', bullet='*')

# --- 11단계 ---
doc.add_heading('[11단계] 7코드(K차크라) 에너지 파동 분석', level=3)
add_rich(doc, '회원이 화면에 나오는 키워드 목록 중 **무의식적으로 끌리는 항목**을 다중 선택합니다.')
add_rich(doc, '**7코드(K차크라):** 선택한 키워드를 기반으로, 7개 에너지 코드 중 **어느 코드의 에너지가 부족하거나 방전되어 있는지** 밸런스를 분석합니다.', bullet='*')
add_rich(doc, '**핵심 포인트:** 에너지가 부족하거나 방전된 코드를 발견하면, 해당 코드의 에너지를 **충전하고 충만하게 관리하는 것**이 핵심입니다. 뇌교육 수련과 일상 습관을 통해 약해진 에너지를 다시 채워 전체 밸런스를 회복합니다.', bullet='*')
add_rich(doc, '**결과 도출:** AI가 **에너지가 가장 부족한(방전된) 코드**를 자동 산출합니다.', bullet='*')
add_rich(doc, '**설명 팁:** "7코드 결과는 점을 보거나 질환을 진단하는 것이 아닙니다. 회원님의 7개 에너지 코드 중 어느 부분이 방전되어 충전이 필요한지 보여주는 심리학적 웰니스 에너지 지표입니다. 부족한 코드를 충전하면 몸과 마음의 에너지가 다시 충만해집니다."', bullet='*')

# ===== Chapter 3 =====
doc.add_heading('Chapter 3. 종합 나이 산출 공식', level=2)

doc.add_heading('3-1. 신체 나이 (Physical Age)', level=3)
add_rich(doc, 'AI가 1~7단계의 사진 + 수치 데이터를 종합 분석하여 산출합니다.')
add_rich(doc, '자세 분석(정면/측면), 균형, 팔 올리기, 유연성: **AI 종합 판단**', indent=0.5)
add_rich(doc, '스쿼트, 팔굽혀펴기: **코드 확정 나이 (AI가 변경 불가)**', indent=0.5)

doc.add_heading('3-2. 두뇌 나이 (Brain Age)', level=3)
add_rich(doc, '8~9단계 점수를 기반으로 **코드가 자동 산출**합니다 (AI 임의 변경 불가).')
add_rich(doc, '**공식:** 두뇌 나이 = (인지반응 나이 + 기억력 나이) / 2', indent=0.5)
add_rich(doc, '**각 테스트별 환산:** 100점 = 20세, 0점 = 90세 (선형)', indent=0.5)
add_rich(doc, '**예시:** 인지반응 80점(34세) + 기억력 60점(48세) = 두뇌 나이 41세', indent=0.5)

doc.add_heading('3-3. 종합 건강 나이 (Comprehensive Age)', level=3)
add_rich(doc, '신체 나이, 두뇌 나이, 얼굴 활력 나이를 통합하여 AI가 최종 산출합니다.')

doc.add_heading('3-4. 종합 점수 (Overall Score, 100점 만점)', level=3)
add_rich(doc, 'AI가 모든 단계의 데이터를 종합하여 **100점 만점 절대평가**로 산출합니다.')

# ===== 푸터 =====
doc.add_paragraph('')
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('---')
run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('(C) 2026 Brain Training Center(BTC) | Confidential')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '맑은 고딕'

# ===== 저장 =====
output_path = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docs\V4_사용_및_평가기준_매뉴얼_v2.docx'
doc.save(output_path)
print(f'[OK] Word file created: {output_path}')
