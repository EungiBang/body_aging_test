import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # ===== 페이지 설정 =====
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===== 스타일 설정 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    def add_rich(text, level=None, bullet=False):
        if level:
            p = doc.add_heading('', level=level)
        else:
            p = doc.add_paragraph()
            if bullet:
                p.style = 'List Bullet'
        
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            if not level:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(11)

    # ===== 표지 생성 =====
    for _ in range(6):
        doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('BTC 3바디 AI 분석기 V4.2')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.name = '맑은 고딕'

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('실전 상담 및 세일즈 매뉴얼')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x8a)
    run.font.name = '맑은 고딕'

    doc.add_paragraph('')
    doc.add_paragraph('')

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_p.add_run('브레인트레이닝센터(BTC)')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2d, 0x3a, 0x8c)
    run.font.name = '맑은 고딕'

    doc.add_page_break()

    # ===== 본문 파싱 =====
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    skip_title = True
    for line in lines:
        line = line.strip()
        if not line:
            if not in_table:
                doc.add_paragraph('')
            continue
            
        # 첫 번째 제목 스킵 (표지로 대체)
        if line.startswith('# ') and skip_title:
            skip_title = False
            continue

        if line.startswith('---'):
            doc.add_page_break()
            continue
            
        if line.startswith('|'):
            in_table = True
            add_rich(line.replace('|', ' | '))
            continue
        in_table = False
            
        if line.startswith('## '):
            add_rich(line[3:], level=2)
        elif line.startswith('### '):
            add_rich(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            add_rich(line[2:], bullet=True)
        else:
            add_rich(line)

    doc.save(docx_path)
    print(f'[OK] Word file created: {docx_path}')

md_file = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docs\V4_실전_상담_및_세일즈_매뉴얼.md'
docx_file = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docs\V4_실전_상담_및_세일즈_매뉴얼_완성본.docx'
convert_md_to_docx(md_file, docx_file)
