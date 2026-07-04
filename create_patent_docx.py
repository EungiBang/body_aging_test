# 특허출원 보완용 기술 설명서를 워드(DOCX) 파일로 변환 및 생성하는 스크립트
import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """테이블 셀 패딩 설정 함수"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_patent_document():
    md_path = r"C:\Users\bange\.gemini\antigravity-ide\brain\317316f9-609d-4bc8-aa8c-2f748dc6721b\patent_materials.md"
    docx_path = r"d:\antigravity_vibecoding\BT 3바디 ai테스트\3바디_AI분석기_특허출원_기술설명서_보완.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} 파일이 존재하지 않습니다.")
        return
        
    doc = docx.Document()
    
    # 기본 스타일 폰트 설정 (맑은 고딕 / Calibri)
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 줄간격 설정
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)

    # 마크다운 파일 읽기
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_quote = False
    in_mermaid = False
    
    print("마크다운 파일 파싱 및 워드 파일 생성 중...")
    
    for line in lines:
        stripped = line.strip()
        
        # 1. Mermaid 다이어그램 건너뛰기 (워드 특성상 텍스트 다이어그램은 스킵하거나 텍스트로 처리)
        if stripped.startswith("```mermaid"):
            in_mermaid = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("[시스템 프로세스 흐름도]")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            continue
            
        if in_mermaid:
            if stripped.startswith("```"):
                in_mermaid = False
                continue
            # 다이어그램 텍스트 흐름 표시
            if stripped and not stripped.startswith("graph"):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                # 화살표나 노드 명칭 파싱
                match = re.search(r'\["?(.*?)"?\]', stripped)
                if match:
                    p.add_run(match.group(1))
                else:
                    p.add_run(stripped)
            continue
            
        # 2. 코드 블록 스킵
        if stripped.startswith("```"):
            continue
            
        # 3. 구분선
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run("____________________________________________________")
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # 4. 헤더 파싱
        if stripped.startswith("# "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # 네이비 톤
            continue
            
        elif stripped.startswith("## "):
            h_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(h_text)
            run.font.size = Pt(15)
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)
            continue
            
        elif stripped.startswith("### "):
            h_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(h_text)
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x4A, 0x76, 0xA8)
            continue
            
        # 5. 불릿 목록 (들여쓰기 수준 분석)
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            
            # 수학식 기호 전처리 ($$ 또는 $ 제거)
            bullet_text = bullet_text.replace("$$", "").replace("$", "")
            
            # 들여쓰기 수준 체크 (원시 라인의 공백 수 기준)
            indent_level = len(line) - len(line.lstrip())
            
            if indent_level >= 4:
                # 2레벨 불릿
                p = doc.add_paragraph(style='List Bullet 2')
            else:
                # 1레벨 불릿
                p = doc.add_paragraph(style='List Bullet')
                
            # 강한 강조(**텍스트**) 파싱 및 폰트 서식 적용
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
            continue
            
        # 6. 순서 있는 목록
        match_num = re.match(r'^(\d+)\.\s(.*)', stripped)
        if match_num:
            num = match_num.group(1)
            text = match_num.group(2).replace("$$", "").replace("$", "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            
            r_num = p.add_run(f"{num}. ")
            r_num.bold = True
            
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
            continue
            
        # 7. 일반 본문 문단
        if stripped:
            # 수학식 라인인 경우
            if stripped.startswith("$$") and stripped.endswith("$$"):
                formula = stripped.replace("$$", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(f" 공식: {formula}")
                run.italic = True
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                continue
                
            p = doc.add_paragraph()
            
            # 강조 처리
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    # 인라인 수식 기호 제거
                    clean_part = part.replace("$$", "").replace("$", "")
                    p.add_run(clean_part)
        else:
            # 빈 줄 처리
            pass

    doc.save(docx_path)
    print(f"성공: 워드 파일이 다음 경로에 저장되었습니다: {docx_path}")

if __name__ == "__main__":
    create_patent_document()
