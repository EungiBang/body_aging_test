# 코드맵AI 종합 소개서 v3 - 서비스(점검+상담+코칭) 중심 재작성
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35
rpr = style.element.get_or_add_rPr()
rpr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="맑은 고딕"/>'))

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
DGRAY = RGBColor(0x37, 0x41, 0x51)
MGRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x05, 0x96, 0x69)

def shade(cell, c):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c}" w:val="clear"/>'))

def tbl_borders(tbl, spec):
    pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    pr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}>{spec}</w:tblBorders>'))

CLEAN = (
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
)

def add_table(doc, hdrs, rows, widths=None, hdr_color="1B2A4A"):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = '맑은 고딕'
        shade(c, hdr_color); c.vertical_alignment = 1
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            p = c.paragraphs[0]; r = p.add_run(str(ct))
            r.font.size = Pt(9); r.font.name = '맑은 고딕'; r.font.color.rgb = DGRAY
            c.vertical_alignment = 1
            if ri % 2 == 1: shade(c, "F3F4F6")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    tbl_borders(t._tbl, CLEAN)
    return t

def h1(doc, txt):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(txt); r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = '맑은 고딕'
    bp = doc.add_paragraph(); bp.paragraph_format.space_after = Pt(8)
    bpPr = bp._element.get_or_add_pPr()
    bpPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="2B6CB0"/></w:pBdr>'))

def h2(doc, txt):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(txt); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'

def body(doc, txt, **kw):
    p = doc.add_paragraph()
    if kw.get('indent'): p.paragraph_format.left_indent = Cm(0.5)
    if kw.get('align'): p.alignment = kw['align']
    r = p.add_run(txt)
    r.font.size = Pt(kw.get('size', 10)); r.font.bold = kw.get('bold', False)
    r.font.color.rgb = kw.get('color', DGRAY); r.font.name = '맑은 고딕'
    return p

def bullet(doc, prefix, txt):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'  {prefix}'); r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'
    r = p.add_run(f' --- {txt}'); r.font.size = Pt(9.5); r.font.color.rgb = DGRAY; r.font.name = '맑은 고딕'

def callout(doc, txt, bg="E8F4FD", border="2B6CB0"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; c.text = ''
    r = c.paragraphs[0].add_run(f'  {txt}')
    r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(int(border[:2],16), int(border[2:4],16), int(border[4:],16)); r.font.name = '맑은 고딕'
    shade(c, bg)
    tbl_borders(t._tbl,
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="{border}"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>')
    doc.add_paragraph()

def quote_box(doc, txt):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'
    shade(c, "F8FAFC")
    tbl_borders(t._tbl,
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="2B6CB0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>')
    doc.add_paragraph()


# ═══════════════ 표지 ═══════════════
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CODEMAP AI'); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = '맑은 고딕'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('코드맵AI'); r.font.size = Pt(24); r.font.color.rgb = ACCENT; r.font.name = '맑은 고딕'

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="2B6CB0"/></w:pBdr>'))
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('3바디 7코드 AI 건강 솔루션'); r.font.size = Pt(16); r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'
doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"몸, 마음, 뇌의 균형을 측정하고\n맞춤형 건강 솔루션을 제공합니다."')
r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = MGRAY; r.font.name = '맑은 고딕'

for _ in range(3): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('수십 년간 축적된 뇌교육 전문 노하우와 AI 기술의 만남'); r.font.size = Pt(11); r.font.color.rgb = DGRAY; r.font.name = '맑은 고딕'
doc.add_paragraph()

for org in ['한국뇌과학연구원  |  국가공인 브레인트레이너협회', '글로벌사이버대학교 뇌교육학과 / 명상치유학과  |  브레인트레이닝센터']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(org); r.font.size = Pt(9.5); r.font.color.rgb = ACCENT; r.font.name = '맑은 고딕'

for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('기업 건강경영 솔루션 제안서  |  2026'); r.font.size = Pt(11); r.font.color.rgb = MGRAY; r.font.name = '맑은 고딕'

doc.add_page_break()


# ═══════════════ 목차 ═══════════════
h1(doc, '목차'); doc.add_paragraph()
toc = [
    ('01', '코드맵AI 소개'),
    ('02', '3바디(3BODY) -- 몸, 마음, 뇌의 통합 건강'),
    ('03', '7코드(7CODE) -- 에너지 밸런스와 맞춤 솔루션'),
    ('04', '연구 기반 및 전문가 네트워크'),
    ('05', 'AI 측정 프로세스'),
    ('06', '서비스 패키지 -- 점검, 상담, 코칭'),
    ('07', '기업 도입 효과'),
    ('08', '도입 절차'),
    ('09', '도입 실적 및 레퍼런스'),
]
for num, title in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'{num}  '); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = ACCENT; r.font.name = '맑은 고딕'
    r = p.add_run(title); r.font.size = Pt(12); r.font.color.rgb = DGRAY; r.font.name = '맑은 고딕'

doc.add_page_break()


# ═══════════════ 01. 코드맵AI 소개 ═══════════════
h1(doc, '01. 코드맵AI 소개')

body(doc, '오늘날 직장인의 건강 문제는 단순한 체력 저하를 넘어, 만성 스트레스, 인지 기능 저하, 감정적 소진(번아웃)이 복합적으로 얽혀 있습니다. 헬스장에서 몸만 단련하거나, 연 1회 건강검진의 숫자만으로는 이 복합적인 건강 문제를 파악하기 어렵습니다.')
doc.add_paragraph()

quote_box(doc, '"진정한 건강이란, 몸(Body)과 마음(Mind)과 뇌(Brain)가\n조화롭게 균형을 이루고 있는 상태입니다."')

body(doc, '코드맵AI(CODEMAP AI)는 브레인트레이닝센터가 수십 년간 축적해 온 몸, 마음, 뇌 통합 건강관리 노하우를 바탕으로, 한국뇌과학연구원, 국가공인 브레인트레이너협회, 글로벌사이버대학교 뇌교육학과 및 명상치유학과의 전문가들이 공동 연구개발한 차세대 AI 건강 솔루션입니다.')
doc.add_paragraph()

body(doc, '코드맵AI는 직원의 신체 능력과 인지 기능, 에너지 밸런스를 종합 측정하여 건강 상태를 3바디(몸/마음/뇌)로 분석하고, 7코드 에너지 체계를 기반으로 각 개인에게 최적화된 건강 솔루션을 제공합니다. 단순히 "어디가 나쁜지"를 알려주는 것이 아니라, "왜 나쁜지, 어떻게 회복할 수 있는지"까지 전문가 상담과 프로그램 코칭을 통해 완결합니다.')

doc.add_paragraph()

h2(doc, '기존 건강 측정 방식과의 차이')

add_table(doc,
    ['구분', '기존 건강검진 / 체력측정', '코드맵AI'],
    [
        ['측정 범위', '신체 수치 중심\n(키, 몸무게, 혈압, 악력 등)', '몸 + 마음 + 뇌\n통합 건강 상태 측정'],
        ['건강 솔루션', '수치 알림에 그침\n(구체적 해결책 부재)', '7코드 기반 맞춤형\n건강 프로그램 제안 + 전문가 상담'],
        ['전문 상담', '없음\n(결과지 배포로 끝)', '브레인트레이닝 전문가의\n1:1 결과 상담 및 코칭'],
        ['소요 시간', '1인당 30~60분', '1인당 5~8분\n(라이트버전 2~3분)'],
        ['전문 인력', '체육 측정사 또는 간호사 필수', '브레인트레이닝전문가\n/ 뇌건강전문가'],
        ['결과물', '단순 수치 나열', 'AI 종합 분석 리포트\n(3바디 진단 + 7코드 솔루션)'],
    ],
    widths=[3, 6, 6]
)

doc.add_page_break()


# ═══════════════ 02. 3바디 ═══════════════
h1(doc, '02. 3바디(3BODY) -- 몸, 마음, 뇌의 통합 건강')

body(doc, '코드맵AI의 건강 철학은 "3바디(3BODY)"에 기반합니다. 사람의 건강을 몸(Body), 마음(Mind), 뇌(Brain)의 세 축으로 바라보며, 이 세 가지가 삼위일체로 균형을 이루어야 진정한 건강이라고 정의합니다.')
doc.add_paragraph()
body(doc, '기존의 건강 측정은 대부분 "몸"의 수치에만 집중합니다. 그러나 현대인의 건강 문제는 스트레스(마음), 인지 기능 저하(뇌), 만성 피로(몸)가 서로 연쇄적으로 영향을 주고받으며 발생합니다. 코드맵AI는 이 세 가지를 모두 측정하고 분석합니다.')
doc.add_paragraph()

add_table(doc,
    ['영역', '의미', '코드맵AI가 측정하는 것', '왜 중요한가'],
    [
        ['BODY\n(몸)', '신체 구조의 균형,\n근력, 유연성', '체형 정렬(비대칭, 거북목),\n하체/상체 근력, 균형 감각,\n유연성, 안면 노화도', '몸의 불균형은 통증과 부상의\n시작점이며, 에너지 흐름을\n방해하는 물리적 원인입니다.'],
        ['MIND\n(마음)', '감정 에너지의 흐름,\n스트레스 상태', '7코드 에너지 밸런스 분석\n(무의식적 에너지 상태 측정)', '감정적 소진(번아웃)과\n만성 스트레스는 신체 노화를\n가속시키는 핵심 요인입니다.'],
        ['BRAIN\n(뇌)', '인지 기능,\n뇌 신경망의 활력', '반응 속도, 억제 기능(전두엽),\n작업 기억력(Working Memory)', '뇌 기능의 저하는 판단력,\n집중력, 업무 효율을\n직접적으로 떨어뜨립니다.'],
    ],
    widths=[2.5, 3.5, 5, 5]
)
doc.add_paragraph()

callout(doc, '코드맵AI는 세 영역의 점수를 종합하여 "3바디 통합 밸런스 나이"를 산출합니다. 이는 단순한 체력 나이가 아니라, 몸과 마음과 뇌의 건강이 모두 반영된 "진정한 건강 나이"입니다.', "F5F3FF", "5B21B6")

doc.add_page_break()


# ═══════════════ 03. 7코드 ═══════════════
h1(doc, '03. 7코드(7CODE) -- 에너지 밸런스와 맞춤 솔루션')

body(doc, '3바디가 "어디가 약한지"를 진단한다면, 7코드는 "왜 약한지, 어떻게 회복할지"를 안내하는 솔루션 체계입니다.')
doc.add_paragraph()
body(doc, '7코드(7CODE)는 인체의 에너지가 흐르는 7개의 핵심 포인트를 의미합니다. 각 코드는 신체의 특정 부위, 감정 상태, 그리고 에너지의 흐름과 연결되어 있습니다. 코드맵AI는 AI가 측정한 신체 데이터를 7코드 체계에 매칭하여, 에너지가 부족하거나 정체된 코드를 정확히 찾아내고, 그 코드를 회복시키는 맞춤형 프로그램을 자동으로 제안합니다.')
doc.add_paragraph()

h2(doc, '7코드 에너지 체계')

add_table(doc,
    ['코드', '에너지 영역', '관련 신체 부위', '부족 시 나타나는 증상'],
    [
        ['1코드\n(기초/생명)', '생명력, 토대,\n그라운딩', '골반, 하체', '만성 피로, 의욕 저하,\n하체 무력감'],
        ['2코드\n(감정/창조)', '유연성, 감정 흐름,\n장 건강', '단전, 허리', '소화 장애, 감정 경직,\n허리 통증'],
        ['3코드\n(의지/추진)', '코어 근력,\n자기표현, 추진력', '복부, 위장', '무기력, 추진력 부족,\n중심 흔들림'],
        ['4코드\n(순환/포용)', '심폐 순환,\n감정 순환', '가슴, 어깨', '답답함, 호흡 얕아짐,\n감정적 피로(번아웃)'],
        ['5코드\n(소통/균형)', '소통, 표현력,\n경추 건강', '목, 어깨', '만성 어깨 결림, 두통,\n소통 어려움'],
        ['6코드\n(집중/통찰)', '뇌 신경망,\n직관, 평형 감각', '뇌, 소뇌', '집중력 저하, 어지러움,\n수면의 질 하락'],
        ['7코드\n(조화/통합)', '전체 조화,\n회복탄력성', '정수리, 전신 통합', '만성 불면, 전반적 무기력,\n삶의 활력 저하'],
    ],
    widths=[2.5, 3, 3, 5]
)
doc.add_paragraph()

h2(doc, '7코드 기반 맞춤 솔루션 제안')
body(doc, 'AI가 측정한 데이터를 기반으로 에너지가 가장 부족한 코드를 자동으로 찾아내어, 해당 코드를 회복시키는 최적의 프로그램을 제안합니다.')
doc.add_paragraph()

add_table(doc,
    ['취약 코드', '권장 프로그램', '기대 효과'],
    [
        ['하위 코드\n(1, 2코드) 부족', '충전명상 수련\n바디프리 / 장청뇌청', '생명력 회복, 하체 안정화,\n장 건강 개선, 유연성 회복'],
        ['중간 코드\n(3, 4코드) 부족', '충전명상 수련 / 클린호흡\n솔라시스템 / 마음프리', '코어 강화, 심폐 순환 개선,\n감정 해소, 호흡 깊어짐'],
        ['상위 코드\n(5, 6, 7코드) 부족', '충전명상 수련\nPBM / 성인운기스쿨', '뇌 신경망 활성화, 수면 질 향상,\n집중력 회복, 전체 밸런스 조화'],
    ],
    widths=[3.5, 5, 7]
)
doc.add_paragraph()

callout(doc, '단순히 "운동을 하세요" "스트레스를 줄이세요"라는 막연한 조언이 아닙니다. AI가 데이터를 기반으로 정확히 어떤 에너지가 부족한지 진단하고, 전문가가 그에 맞는 구체적인 프로그램을 직접 코칭합니다. 이것이 코드맵AI가 기존 건강 솔루션과 근본적으로 다른 점입니다.', "F5F3FF", "5B21B6")

doc.add_page_break()


# ═══════════════ 04. 연구 기반 ═══════════════
h1(doc, '04. 연구 기반 및 전문가 네트워크')

body(doc, '코드맵AI는 단순한 IT 기술이 아닙니다. 수십 년간 대한민국 뇌훈련과 뇌교육 분야를 이끌어 온 전문가들의 연구와 현장 경험이 AI 기술과 결합된 솔루션입니다.')
doc.add_paragraph()

h2(doc, '공동 연구개발 기관')
add_table(doc,
    ['기관명', '역할 및 전문 분야'],
    [
        ['브레인트레이닝센터', '전국 운영 센터 네트워크를 통해 축적한 수십 년간의 몸, 마음, 뇌 통합 건강관리 노하우와 현장 데이터를 기반으로 한 3바디 7코드 체계의 원천 개발'],
        ['한국뇌과학연구원', '뇌과학 연구 데이터를 기반으로 한 인지 기능 평가 모델 설계, 뇌 건강 지표의 과학적 근거 제공'],
        ['국가공인\n브레인트레이너협회', '국가공인 브레인트레이너 자격 교육과정에서 검증된 건강 평가 기준 및 솔루션 프로그램 체계 자문'],
        ['글로벌사이버대학교\n뇌교육학과\n/ 명상치유학과', '학술적 연구를 통한 3바디 7코드 체계의 이론적 검증, 에너지 밸런스 측정 모델의 학문적 기반 구축'],
    ],
    widths=[4, 12]
)
doc.add_paragraph()

h2(doc, '적용 기술 및 참조 프레임워크')
add_table(doc,
    ['영역', '적용 내용', '설명'],
    [
        ['신체 자세 분석', 'Kendall 자세 평가법', '세계적으로 인정받는 자세 정렬 분석의 표준 기준'],
        ['체력 기준', '국민체력100 기준 참조', '대한민국 국민체력인증센터의 성별/연령별 체력 데이터'],
        ['인지 기능 평가', 'Stroop 인지 과제', '전두엽 억제 기능 측정의 표준 신경심리학 검사'],
        ['AI 자세 추정', 'Google TensorFlow.js', '실시간 33개 관절 좌표 추출 (초당 30프레임)'],
        ['AI 종합 분석', 'Google Gemini AI', '촬영 이미지와 측정 데이터를 종합 분석하는 생성형 AI'],
        ['에너지 체계', '3바디 7코드', '브레인트레이닝센터 고유의 통합 에너지 밸런스 체계'],
    ],
    widths=[3, 4.5, 8.5]
)
doc.add_paragraph()
callout(doc, '코드맵AI는 현장의 전문가가 수천 명의 회원을 직접 만나며 쌓아 올린 실전 노하우, 뇌과학 연구자들의 학술적 검증, 그리고 Google AI의 첨단 기술이 삼위일체로 결합된 솔루션입니다.')

doc.add_page_break()


# ═══════════════ 05. AI 측정 프로세스 ═══════════════
h1(doc, '05. AI 측정 프로세스')

body(doc, '코드맵AI의 측정은 카메라 앞에서 AI의 음성 안내를 따라 자연스럽게 진행됩니다. 브레인트레이닝 전문가들이 안내를 하여 전문적으로 점검 받아볼 수 있습니다.')
doc.add_paragraph()

h2(doc, '신체(BODY) 측정 -- 8단계 AI 가이드 진단 (라이트버전 6단계 압축)')
add_table(doc,
    ['단계', '측정 항목', '방식', 'AI가 분석하는 것'],
    [
        ['1', '정면 체형', '5초 정지 촬영', '어깨 비대칭, 골반 틀어짐, 무릎 정렬'],
        ['2', '측면 체형', '5초 정지 촬영', '거북목 각도, 전방 경사, 척추 만곡'],
        ['3', '균형 감각\n(한발 서기)', '15초\n실시간 추적', '발 디딤 횟수, 상체 흔들림 지수'],
        ['4', '상지 가동범위', '순간 캡처', '어깨 충돌 구간, 양측 높이 차이'],
        ['5', '유연성 (전굴)', '순간 캡처', '손끝-바닥 거리, 햄스트링 이완도'],
        ['6', '하체 근력\n(스쿼트)', '15초\n자동 카운팅', '횟수 + 자세 정확도'],
        ['7', '상체 근력\n(푸시업)', '15초\n자동 카운팅', '횟수 + 자세 정확도'],
        ['8', '안면 노화 분석', '안면 촬영', '피부 탄력, 주름 분포, 노화도'],
    ],
    widths=[1.5, 3.5, 3, 7.5]
)
doc.add_paragraph()

h2(doc, '뇌(BRAIN) 측정 -- 인지 기능 3단계 평가')
add_table(doc,
    ['단계', '측정 항목', '평가 내용'],
    [
        ['1', '반응 속도 테스트', '시각 자극에 대한 밀리초 단위 반응 시간 측정'],
        ['2', '억제 기능 (스트룹 변형)', '방해 자극 속 정답률, 충동 억제 능력 평가'],
        ['3', '작업 기억력 테스트', '다중 항목 기억 및 회상 정확도 측정'],
    ],
    widths=[1.5, 4, 10]
)
doc.add_paragraph()

h2(doc, '마음(MIND) 측정 -- 7코드 에너지 밸런스 분석')
body(doc, 'AI가 측정한 신체 데이터(자세 정렬, 근력 분포, 유연성, 균형 감각 등)를 7코드 에너지 체계에 자동 매칭하여, 어떤 에너지가 부족하고 어디에서 정체가 일어나고 있는지를 진단합니다.')
doc.add_paragraph()

h2(doc, 'AI 종합 리포트 산출')
body(doc, '모든 측정이 완료되면 AI가 자동으로 종합 분석 리포트를 생성합니다.')
bullet(doc, '개별 채점', '각 항목별 100점 만점 절대 점수 산출')
bullet(doc, '종합 점수', '6개 신체 항목의 가중치 합산으로 종합 점수 산출')
bullet(doc, '3바디 건강 나이', '종합 점수를 기반으로 실제 나이 대비 신체 나이 자동 계산')
bullet(doc, '7코드 솔루션', '에너지가 부족한 코드를 찾아 맞춤 프로그램 자동 제안')

doc.add_page_break()


# ═══════════════ 06. 서비스 패키지 (핵심 신규 섹션) ═══════════════
h1(doc, '06. 서비스 패키지 -- 점검, 상담, 코칭')

body(doc, '코드맵AI는 단순히 장비나 소프트웨어를 판매하는 것이 아닙니다. AI 점검부터 전문가 상담, 맞춤 프로그램 코칭까지 하나의 서비스로 제공하여, 기업의 직원 건강관리를 완결합니다.')
doc.add_paragraph()
body(doc, '기업의 목적과 규모에 따라 3가지 서비스 패키지를 선택할 수 있습니다.')
doc.add_paragraph()

# ── 패키지 1 ──
h2(doc, 'BASIC  |  코드맵 점검 + 결과 안내')

body(doc, '코드맵AI를 통해 직원의 몸, 마음, 뇌 건강 상태를 측정하고, AI가 생성한 종합 리포트를 전달합니다.', indent=True)
doc.add_paragraph()

add_table(doc,
    ['항목', '내용'],
    [
        ['서비스 내용', 'AI 건강 점검 (3바디 7코드 진단)\nAI 종합 리포트 제공 (개인별 결과지)\n7코드 에너지 밸런스 결과 안내'],
        ['운영 방식', '행사장, 세미나실, 회의실 등에 장비 세팅\n브레인트레이닝 전문가 파견 운영'],
        ['적합한 상황', '건강의 날 행사, 채용 박람회, 복지 이벤트 등\n대규모 인원의 건강 체험 행사'],
        ['직원 소요 시간', '1인당 약 5~8분 (라이트버전 2~3분)'],
        ['기업이 얻는 것', '직원들의 건강 인식 제고\n첨단 AI 체험형 복지 이벤트로 만족도 향상\n전체 참여자 건강 통계 리포트 (선택)'],
    ],
    widths=[3, 13]
)
doc.add_paragraph()

# ── 패키지 2 ──
h2(doc, 'STANDARD  |  코드맵 점검 + 1:1 전문가 상담')

body(doc, 'AI 점검에 더하여, 브레인트레이닝 전문가가 직원 개인별로 결과를 해석해 주고, 건강 상태의 원인과 개선 방향을 1:1로 상담합니다.', indent=True)
doc.add_paragraph()

add_table(doc,
    ['항목', '내용'],
    [
        ['서비스 내용', 'BASIC 패키지의 모든 내용 포함\n+ 브레인트레이닝 전문가의 1:1 결과 상담\n+ 3바디 균형 상태 해석 및 원인 분석\n+ 7코드 취약 코드 기반 개선 방향 안내'],
        ['운영 방식', '건강관리실, 상담실 등 전용 공간 세팅\n전문가 파견 (상담 인원에 따라 조정)'],
        ['적합한 상황', '분기별/반기별 직원 건강관리 프로그램\n근골격계 부담작업 관리 사전 점검\n번아웃 위험군 선별 및 조기 대응'],
        ['직원 소요 시간', '1인당 약 15~20분 (점검 + 상담)'],
        ['기업이 얻는 것', '직원 개개인의 건강 상태에 대한 전문적 해석 제공\n자발적 건강 행동 변화 유도 (전문가 상담 효과)\n부서별/연령별 건강 리스크 분석 리포트'],
    ],
    widths=[3, 13]
)
doc.add_paragraph()

# ── 패키지 3 ──
h2(doc, 'PREMIUM  |  코드맵 점검 + 1:1 상담 + 프로그램 코칭')

body(doc, '점검과 상담을 넘어, 직원들이 실제로 건강을 회복할 수 있도록 7코드 기반의 맞춤형 프로그램을 직접 코칭합니다. 개별, 소규모 그룹, 단체 등 기업의 여건에 맞는 형태로 운영됩니다.', indent=True)
doc.add_paragraph()

add_table(doc,
    ['항목', '내용'],
    [
        ['서비스 내용', 'STANDARD 패키지의 모든 내용 포함\n+ 7코드 취약 코드 회복을 위한 맞춤 프로그램 코칭\n+ 충전명상 수련, 바디프리, 클린호흡, PBM 등\n+ 개별 / 소규모 그룹 / 단체 세션 운영'],
        ['운영 방식', '기업 내 강당, 세미나실, 휴게 공간 등\n전문 코칭 트레이너 파견'],
        ['운영 형태', '개별 코칭: 1:1 맞춤형 프로그램 (고위 임원, VIP)\n소규모 그룹: 4~8명 단위 (팀별/부서별)\n단체 세션: 20명 이상 (전사 건강 프로그램)'],
        ['적합한 상황', '직원 번아웃 예방 및 회복 프로그램\n리더십/임원 대상 프리미엄 건강관리\n팀 빌딩과 연계한 건강 워크숍\n장기적 직원 건강경영 시스템 구축'],
        ['직원 소요 시간', '점검+상담: 약 15~20분\n프로그램 코칭: 60~90분 (세션당)'],
        ['기업이 얻는 것', '측정에서 개선까지 원스톱 건강관리 실현\n직원의 실질적 건강 회복 및 활력 향상\n조직 생산성 향상 및 이직률 감소 기여\n건강경영 우수기업 인증 근거 자료 확보'],
    ],
    widths=[3, 13]
)

doc.add_paragraph()

callout(doc, '모든 패키지는 기업의 규모, 예산, 목적에 맞춰 유연하게 설계됩니다. 단발성 행사부터 연간 정기 프로그램까지, 귀사에 최적화된 건강경영 파트너가 되겠습니다.', "ECFDF5", "059669")

doc.add_page_break()


# ═══════════════ 07. 기업 도입 효과 ═══════════════
h1(doc, '07. 기업 도입 효과')

h2(doc, '직원이 체감하는 효과')
add_table(doc,
    ['효과', '상세 내용'],
    [
        ['"내 몸 상태를\n눈으로 확인"', 'AI가 측정한 체형 사진 위에 비대칭, 거북목, 균형 지수 등이 시각적으로 표시됩니다.\n숫자만 나열하는 건강검진과 달리, 자신의 상태를 직관적으로 이해할 수 있습니다.'],
        ['"막연한 불안이\n구체적 행동으로"', '7코드 솔루션이 "어떤 프로그램을, 얼마나 해야 하는지"를 구체적으로 안내합니다.\n전문가 상담을 통해 내 상태의 원인을 이해하고, 실천 가능한 맞춤 가이드를 받습니다.'],
        ['"5분 만에 끝나는\n건강 체크"', '1인당 5~8분이면 측정이 완료됩니다. 업무 시간 중에도 부담 없이 참여할 수 있어\n참여율이 기존 체력 측정 대비 크게 향상됩니다.'],
        ['"재미있는\nAI 체험"', 'AI 음성 안내, 실시간 자세 추적, 인지 게임 등 체험형 요소가 풍부하여\n직원들의 자발적 참여와 긍정적 반응을 이끌어냅니다.'],
        ['"전문가에게\n직접 상담"', '결과지만 받아보는 것이 아니라, 브레인트레이닝 전문가가 내 결과를 해석하고\n개선 방법을 직접 안내합니다. 건강에 대한 동기부여가 완전히 달라집니다.'],
    ],
    widths=[3.5, 12.5]
)
doc.add_paragraph()

h2(doc, '기업이 얻는 효과')
add_table(doc,
    ['효과', '상세 내용'],
    [
        ['복지 만족도 향상', '첨단 AI 체험형 건강 복지 프로그램으로 직원 만족도와 기업 이미지 동시 향상'],
        ['번아웃 예방', '3바디 7코드 분석으로 스트레스 및 번아웃 위험군을 조기 발견,\n전문가 상담과 프로그램으로 사전 개입 가능'],
        ['근골격계 질환 예방', '자세 비대칭, 근력 부족 등 근골격계 위험군을 사전에 파악\n산업안전보건법상 근골격계 부담작업 관리에 활용 가능'],
        ['조직 생산성 향상', '직원의 신체 활력과 인지 기능 회복으로 업무 집중력 및 효율 향상\n건강한 조직 문화 형성 기여'],
        ['데이터 기반\n건강 경영', '참여자 전체의 건강 통계 리포트 제공\nEAP(직원지원프로그램) 등 복지 정책 수립의 객관적 근거 확보'],
        ['건강경영 인증 지원', '건강경영 우수기업 인증(고용노동부) 등\n대외 인증 획득 시 활용 가능한 프로그램 운영 실적 확보'],
    ],
    widths=[3.5, 12.5]
)

doc.add_page_break()


# ═══════════════ 08. 도입 절차 ═══════════════
h1(doc, '08. 도입 절차')

body(doc, '코드맵AI 서비스는 상담부터 실행까지 간결하고 유연하게 진행됩니다.')
doc.add_paragraph()

add_table(doc,
    ['단계', '내용', '소요', '상세'],
    [
        ['STEP 1', '상담 및\n데모 체험', '1~2일', '방문 또는 온라인으로 코드맵AI를 직접 체험해 보실 수 있습니다.\n기업의 목적과 규모에 맞는 패키지를 함께 설계합니다.'],
        ['STEP 2', '서비스 설계', '3~5일', '참여 인원, 일정, 장소, 운영 형태(개별/그룹/단체) 확정\n리포트 양식, 로고 등 커스터마이징 요소 협의'],
        ['STEP 3', '현장 세팅\n및 준비', '당일', '장비(카메라, PC/디스플레이) 설치 및 테스트\n전문가 파견 및 운영 동선 확인'],
        ['STEP 4', '서비스 실행', '당일~\n기간별', '코드맵 점검 + 결과 안내 / 1:1 상담 / 프로그램 코칭\n선택한 패키지에 따라 운영'],
        ['STEP 5', '결과 보고', '서비스 후\n3~5일', '참여자 전체 건강 통계 리포트 제공 (선택)\n향후 정기 프로그램 운영 방안 제안'],
    ],
    widths=[2.5, 3, 2.5, 8]
)
doc.add_paragraph()

callout(doc, '단발성 건강 행사부터 분기별 정기 프로그램, 연간 건강경영 파트너십까지 기업의 니즈에 맞춰 유연하게 설계됩니다. 부담 없이 상담을 요청해 주세요.', "FEF3C7", "D97706")

doc.add_page_break()


# ═══════════════ 09. 도입 실적 ═══════════════
h1(doc, '09. 도입 실적 및 레퍼런스')

add_table(doc,
    ['구분', '내용'],
    [
        ['도입 기관', '브레인트레이닝센터 (전국 다수 지점)'],
        ['운영 방식', '상시 운영 + 야외 행사 운영 병행'],
        ['시스템 버전', 'v5.1 (2026년 최신)'],
        ['핵심 성과', '코드맵AI 체험형 건강 점검을 통한\n참여율 및 만족도 향상, 프로그램 등록 전환율 증가'],
    ],
    widths=[4, 12]
)
doc.add_paragraph()

h2(doc, '코드맵AI를 선택해야 하는 이유')

reasons = [
    ('몸, 마음, 뇌를 함께 봅니다', '기존 건강 측정이 놓치는 마음의 에너지 상태와 뇌 기능까지 통합 측정합니다.'),
    ('측정에서 솔루션까지, 한 번에', '"어디가 나쁘다"로 끝나지 않습니다. 7코드 체계로 원인을 진단하고, 전문가가 직접 솔루션을 코칭합니다.'),
    ('수십 년의 전문성이 담겨 있습니다', '뇌교육 전문가, 뇌과학 연구자, 현장 전문가들의 노하우가 AI와 서비스 전체에 녹아 있습니다.'),
    ('점검부터 코칭까지, 전문가가 함께합니다', '장비만 보내는 것이 아닙니다. 브레인트레이닝 전문가가 직접 방문하여 점검, 상담, 코칭을 제공합니다.'),
    ('기업의 건강경영을 함께 설계합니다', '단발성 행사부터 연간 프로그램까지, 귀사의 건강경영 파트너로서 함께 성장합니다.'),
]

for title, desc in reasons:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'  {title}'); r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(desc); r.font.size = Pt(10); r.font.color.rgb = DGRAY; r.font.name = '맑은 고딕'


# ═══════════════ 문의 페이지 ═══════════════
doc.add_page_break()
for _ in range(5): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('문의 및 데모 신청'); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name = '맑은 고딕'
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="2B6CB0"/></w:pBdr>'))
doc.add_paragraph()

ct = doc.add_table(rows=6, cols=2); ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (l, v) in enumerate([('솔루션명','코드맵AI (CODEMAP AI)'),('개발사','(회사명 기입)'),('담당자','(담당자명 기입)'),('연락처','(전화번호 기입)'),('이메일','(이메일 기입)'),('웹사이트','(URL 기입)')]):
    cl = ct.rows[i].cells[0]; cl.text = ''; cl.width = Cm(4)
    p = cl.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(l); r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = DARK_BLUE; r.font.name = '맑은 고딕'
    cr = ct.rows[i].cells[1]; cr.text = ''; cr.width = Cm(10)
    r = cr.paragraphs[0].add_run(f'  {v}'); r.font.size = Pt(11); r.font.color.rgb = DGRAY; r.font.name = '맑은 고딕'

tbl_borders(ct._tbl,
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>')

doc.add_paragraph(); doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('현장 방문 또는 온라인 화상으로 코드맵AI를 직접 체험해 보실 수 있습니다.'); r.font.size = Pt(10.5); r.font.color.rgb = MGRAY; r.font.name = '맑은 고딕'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('귀사의 환경에 맞는 최적의 건강경영 솔루션을 함께 설계해 드리겠습니다.'); r.font.size = Pt(10.5); r.font.color.rgb = MGRAY; r.font.name = '맑은 고딕'

for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('(C) 2026 CODEMAP AI. All rights reserved.'); r.font.size = Pt(8); r.font.color.rgb = MGRAY; r.font.name = '맑은 고딕'


# ═══════════════ 저장 ═══════════════
output_path = os.path.join(os.path.dirname(__file__), 'CODEMAP_AI_소개서.docx')
doc.save(output_path)
print(f'[OK] Document created: {output_path}')
