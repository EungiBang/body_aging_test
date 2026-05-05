import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# ===== 페이지 설정 =====
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== 스타일 설정 =====
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

def add_rich(doc, text, indent=0, bullet=''):
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    if bullet:
        run = p.add_run(bullet + ' ')
        run.font.name = '맑은 고딕'
        
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
    return p

# ===== 제목 =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('🚀 BTC 3바디 AI진단 V4 지점 설치 가이드')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = '맑은 고딕'

doc.add_paragraph('')
add_rich(doc, '원장님, 새로운 **BT 3바디 AI진단 V4.2.2** 버전을 지점 PC에 설치하는 방법입니다.')
add_rich(doc, '컴퓨터를 잘 모르시는 분도 아래 순서대로만 따라 하시면 5분 안에 쉽게 설치하실 수 있습니다!')
doc.add_paragraph('---')

doc.add_heading('📌 1단계: 기존 프로그램 종료하기', level=2)
add_rich(doc, '새로운 버전을 설치하기 전에, 혹시 켜져 있는 기존 3바디 프로그램이 있다면 우측 상단의 X 버튼을 눌러 **완전히 종료**해 주세요.')

doc.add_heading('📌 2단계: 설치 파일 다운로드 및 실행', level=2)
add_rich(doc, '전달받으신 다운로드 링크를 클릭하여 **BT 3바디 AI진단 V4.2.2 Setup.exe** 파일을 다운로드합니다.', bullet='1.')
add_rich(doc, '다운로드가 완료되면 해당 파일을 **더블 클릭하여 실행**합니다.', bullet='2.')

add_rich(doc, '💡 **잠깐! 파란색 보안 경고 창이 떴나요?**')
add_rich(doc, '새로 만든 프로그램이라 윈도우에서 안전을 위해 일시적으로 띄우는 창입니다. 바이러스가 아니니 안심하세요!', indent=0.5)
add_rich(doc, '파란 창에서 **[추가 정보]** 글씨를 클릭합니다.', indent=1, bullet='1.')
add_rich(doc, '오른쪽 아래에 새로 생긴 **[실행]** 버튼을 클릭하시면 설치가 진행됩니다.', indent=1, bullet='2.')

doc.add_heading('📌 3단계: 시스템 권한 배포 코드 인증 (최초 1회)', level=2)
add_rich(doc, '설치가 완료되고 프로그램이 켜지면, 본사 중앙 서버와 연결하기 위한 인증 화면이 나옵니다.')
add_rich(doc, '**지역 선택 & 지점 선택**: 화면에 나오는 목록에서 우리 지점이 속한 지역과 지점명을 선택합니다.', bullet='1.')
add_rich(doc, '**책임 관리자 (원장님)**: 원장님 성함을 입력합니다. (예: 홍길동)', bullet='2.')
add_rich(doc, '**연락처**: 원장님 휴대폰 번호를 입력합니다.', bullet='3.')
add_rich(doc, '**시스템 권한 배포 코드**: 본사에서 미리 전달해 드린 보안 인증 코드를 입력합니다.', bullet='4.')
add_rich(doc, '하단의 **[시스템 권한 요청]** 버튼을 누릅니다. (인증이 완료되면 화면이 자동으로 새로고침 됩니다.)', bullet='5.')

doc.add_heading('📌 4단계: 기존 회원 정보(V3) 불러오기 안내', level=2)
add_rich(doc, '**V4 버전부터는 기존과 완전히 다른 새롭고 안정적인 회원 DB를 사용합니다.**')
add_rich(doc, '처음 설치하시면 회원 목록이 텅 비어있는 것이 정상입니다! 당황하지 마세요.', indent=0.5, bullet='-')
add_rich(doc, '기존 V3 버전에서 등록했던 회원 정보를 한 번에 불러올 수 있는 기능이 **[회원 관리]** 메뉴 내에 준비되어 있습니다. 해당 버튼을 눌러 기존 회원들을 연동해주세요.', indent=0.5, bullet='-')

doc.add_heading('📌 5단계: 카메라 권한 허용 (필수)', level=2)
add_rich(doc, '프로그램에서 회원의 체형을 스캔하기 위해 카메라를 켭니다.')
add_rich(doc, '윈도우나 백신 프로그램에서 **"카메라 접근을 허용하시겠습니까?"** 라는 팝업이 뜨면 반드시 **[허용]** 또는 **[예]**를 눌러주세요.', indent=0.5, bullet='-')
add_rich(doc, '만약 카메라가 까맣게 나오거나 켜지지 않으면, PC에 웹캠이 잘 꽂혀 있는지 확인해 주세요.', indent=0.5, bullet='-')

doc.add_paragraph('---')
doc.add_heading('🎉 설치 완료!', level=3)
add_rich(doc, '이제 바탕화면에 생긴 **[BT 3바디 AI진단 V4]** 아이콘을 더블클릭하여 언제든 프로그램을 시작하실 수 있습니다.')
add_rich(doc, '회원님들께 한 차원 더 높은 수준의 3바디 AI 진단 서비스를 제공해 보세요!')

output_path = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docs\V4_지점_설치_가이드.docx'
doc.save(output_path)
print(f'[OK] Word file created: {output_path}')
