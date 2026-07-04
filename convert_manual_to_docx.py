# -*- coding: utf-8 -*-
# 마크다운 형식의 사용 매뉴얼 문서를 읽어 python-docx를 통해 서식이 적용된 워드 문서로 변환하는 스크립트

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
    # </w:tblBorders>
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def apply_font_to_run(run, font_name='맑은 고딕', size_pt=10, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_paragraph_with_runs(doc, text, style_name=None, font_name='맑은 고딕', size_pt=10.5, color_rgb=None, space_after=6, space_before=0, line_spacing=1.15, left_indent=None, first_line_indent=None):
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    
    # [설명](#링크) 같은 마크다운 링크에서 텍스트만 추출
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    
    # 마크다운 텍스트 내의 bold(**...**)와 inline code(`...`)를 파싱
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            inner_text = token[2:-2]
            run = p.add_run(inner_text)
            apply_font_to_run(run, font_name=font_name, size_pt=size_pt, bold=True, color_rgb=color_rgb)
        elif token.startswith('`') and token.endswith('`'):
            inner_text = token[1:-1]
            run = p.add_run(inner_text)
            apply_font_to_run(run, font_name='Consolas', size_pt=size_pt-0.5, color_rgb=RGBColor(199, 37, 78))
        else:
            run = p.add_run(token)
            apply_font_to_run(run, font_name=font_name, size_pt=size_pt, color_rgb=color_rgb)
    return p

def add_code_block(doc, text_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    
    text = '\n'.join(text_lines)
    run = p.add_run(text)
    apply_font_to_run(run, font_name='Consolas', size_pt=9.5, color_rgb=RGBColor(51, 51, 51))

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F9F9F9"/>')
    pPr.append(shd)
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="5B9BD5"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            inner_text = token[2:-2]
            run = p.add_run(inner_text)
            apply_font_to_run(run, font_name='맑은 고딕', size_pt=10, bold=True, color_rgb=RGBColor(60, 60, 60))
        elif token.startswith('`') and token.endswith('`'):
            inner_text = token[1:-1]
            run = p.add_run(inner_text)
            apply_font_to_run(run, font_name='Consolas', size_pt=9.5, color_rgb=RGBColor(199, 37, 78))
        else:
            run = p.add_run(token)
            apply_font_to_run(run, font_name='맑은 고딕', size_pt=10, color_rgb=RGBColor(80, 80, 80), italic=True)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="DDDDDD"/></w:pBdr>')
    pPr.append(pBdr)

def add_table_from_markdown(doc, table_data):
    rows_filtered = []
    for row in table_data:
        is_separator = True
        for cell in row:
            if cell.strip() and not re.match(r'^:?-+:?$', cell.strip()):
                is_separator = False
                break
        if not is_separator:
            rows_filtered.append(row)
            
    if not rows_filtered:
        return
        
    num_rows = len(rows_filtered)
    num_cols = max(len(row) for row in rows_filtered)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "CCCCCC")
    
    col_width = Inches(6.0 / num_cols)
    
    for r_idx, row_content in enumerate(rows_filtered):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        
        trPr = row._tr.get_or_add_trPr()
        if is_header:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for c_idx, cell_text in enumerate(row_content):
            if c_idx >= num_cols:
                break
            cell = row.cells[c_idx]
            cell.width = col_width
            
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            if is_header:
                set_cell_background(cell, "4F81BD")
            elif r_idx % 2 == 1:
                set_cell_background(cell, "F2F5F8")
                
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            text = cell_text.strip()
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for token in tokens:
                if not token:
                    continue
                if token.startswith('**') and token.endswith('**'):
                    inner_text = token[2:-2]
                    run = p.add_run(inner_text)
                    if is_header:
                        apply_font_to_run(run, font_name='맑은 고딕', size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
                    else:
                        apply_font_to_run(run, font_name='맑은 고딕', size_pt=9.5, bold=True)
                elif token.startswith('`') and token.endswith('`'):
                    inner_text = token[1:-1]
                    run = p.add_run(inner_text)
                    apply_font_to_run(run, font_name='Consolas', size_pt=9.0, color_rgb=RGBColor(199, 37, 78))
                else:
                    run = p.add_run(token)
                    if is_header:
                        apply_font_to_run(run, font_name='맑은 고딕', size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
                    else:
                        apply_font_to_run(run, font_name='맑은 고딕', size_pt=9.5)
                        
def main():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    md_path = r"d:\antigravity_vibecoding\BT_3Body_Outdoor_Lite\docs\BTC_CODEMAP_V2_사용매뉴얼.md"
    docx_path = r"d:\antigravity_vibecoding\BT_3Body_Outdoor_Lite\docs\BTC_CODEMAP_V2_사용매뉴얼.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_block_lines = []
    
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 코드 블록 처리
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, code_block_lines)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_block_lines.append(line.rstrip('\r\n'))
            i += 1
            continue
            
        # 테이블 처리
        if stripped.startswith("|"):
            in_table = True
            cells = [c.strip() for c in stripped.split('|')]
            if cells and cells[0] == '':
                cells.pop(0)
            if cells and cells[-1] == '':
                cells.pop()
            table_lines.append(cells)
            i += 1
            continue
        else:
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
                
        # 빈 줄 처리
        if not stripped:
            i += 1
            continue
            
        # blockquote 처리
        if stripped.startswith(">"):
            text = stripped[1:].strip()
            add_blockquote(doc, text)
            i += 1
            continue
            
        # 수평선
        if stripped == "---":
            add_horizontal_line(doc)
            i += 1
            continue
            
        # 헤더 처리
        if stripped.startswith("#"):
            header_match = re.match(r'^(#+)\s+(.*)$', stripped)
            if header_match:
                level = len(header_match.group(1))
                title = header_match.group(2)
                
                if level == 1:
                    size = 18
                    color = RGBColor(47, 85, 151)
                    space_before = 18
                    space_after = 8
                elif level == 2:
                    size = 14
                    color = RGBColor(47, 85, 151)
                    space_before = 14
                    space_after = 6
                elif level == 3:
                    size = 12
                    color = RGBColor(60, 60, 60)
                    space_before = 10
                    space_after = 4
                else:
                    size = 11
                    color = RGBColor(80, 80, 80)
                    space_before = 8
                    space_after = 4
                    
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(title)
                apply_font_to_run(run, font_name='맑은 고딕', size_pt=size, bold=True, color_rgb=color)
                
                i += 1
                continue
                
        # 리스트 처리 (글머리 기호 및 번호 목록)
        leading_spaces = len(line) - len(line.lstrip(' \t'))
        indent_level = leading_spaces // 2
        
        bullet_match = re.match(r'^[\-\*\+]\s+(.*)$', stripped)
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        
        left_indent = Inches(0.25 * (indent_level + 1))
        first_line_indent = Inches(-0.15)
        
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = left_indent
            p.paragraph_format.first_line_indent = first_line_indent
            
            bullet_char = '•' if indent_level == 0 else ('◦' if indent_level == 1 else '▪')
            run_bullet = p.add_run(f"{bullet_char}\t")
            apply_font_to_run(run_bullet, font_name='맑은 고딕', size_pt=10.5)
            
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for token in tokens:
                if not token:
                    continue
                if token.startswith('**') and token.endswith('**'):
                    inner_text = token[2:-2]
                    run = p.add_run(inner_text)
                    apply_font_to_run(run, font_name='맑은 고딕', size_pt=10.5, bold=True)
                elif token.startswith('`') and token.endswith('`'):
                    inner_text = token[1:-1]
                    run = p.add_run(inner_text)
                    apply_font_to_run(run, font_name='Consolas', size_pt=10.0, color_rgb=RGBColor(199, 37, 78))
                else:
                    run = p.add_run(token)
                    apply_font_to_run(run, font_name='맑은 고딕', size_pt=10.5)
                    
            i += 1
            continue
            
        elif num_match:
            num_str = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = left_indent
            p.paragraph_format.first_line_indent = first_line_indent
            
            run_num = p.add_run(f"{num_str}.\t")
            apply_font_to_run(run_num, font_name='맑은 고딕', size_pt=10.5)
            
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for token in tokens:
                if not token:
                    continue
                if token.startswith('**') and token.endswith('**'):
                    inner_text = token[2:-2]
                    run = p.add_run(inner_text)
                    apply_font_to_run(run, font_name='맑은 고딕', size_pt=10.5, bold=True)
                elif token.startswith('`') and token.endswith('`'):
                    inner_text = token[1:-1]
                    run = p.add_run(inner_text)
                    apply_font_to_run(run, font_name='Consolas', size_pt=10.0, color_rgb=RGBColor(199, 37, 78))
                else:
                    run = p.add_run(token)
                    apply_font_to_run(run, font_name='맑은 고딕', size_pt=10.5)
                    
            i += 1
            continue
            
        # 일반 본문 텍스트 처리
        if stripped.startswith("!["):
            i += 1
            continue
            
        left_indent_val = Inches(0.25 * indent_level) if indent_level > 0 else None
        add_paragraph_with_runs(doc, stripped, space_after=6, space_before=0, left_indent=left_indent_val)
        i += 1
        
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)
        
    doc.save(docx_path)
    print(f"Word document saved to: {docx_path}")

if __name__ == "__main__":
    main()
