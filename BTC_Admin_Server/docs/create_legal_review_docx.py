# -*- coding: utf-8 -*-
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 기본 페이지 설정 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 스타일 헬퍼 ──
def set_run_style(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._r.get_or_add_rPr()
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: (0,51,102), 2: (0,84,166), 3: (60,60,60)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_style(run, size=sizes[level], bold=True, color=colors[level])
    if level == 1:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(doc, text, indent=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_style(run, size=10.5, color=color)
    return p

def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

def add_divider(doc):
    p = doc.add_paragraph('─' * 55)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(180,180,180)
    run.font.size = Pt(9)

# ══════════════════════════════════════════
# 표지
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = title_p.add_run('BTC 3바디 AI분석기 v1.8.4\n의료법 저촉 사항 검토 보고서')
r1.font.size = Pt(20)
r1.font.bold = True
r1.font.name = '맑은 고딕'
r1.font.color.rgb = RGBColor(0, 51, 102)

add_blank(doc)
add_divider(doc)
add_blank(doc)

today = datetime.date.today().strftime('%Y년 %m월 %d일')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'작성일: {today}   |   버전: v1.8.4   |   분류: 대외비')
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(90,90,90)

add_blank(doc)
doc.add_page_break()

# ══════════════════════════════════════════
# 1. 검토 근거 법령
# ══════════════════════════════════════════
add_heading(doc, '1. 검토 근거 법령', level=2)
laws = [
    ('의료법 제2조', '의료인·의료행위 정의 — 진찰·검안·처방 등은 면허 의료인만 수행 가능'),
    ('의료법 제27조', '무면허 의료행위 금지 — 의료인이 아닌 자는 의료행위 불가 (위반 시 3년 이하 징역 또는 3천만 원 이하 벌금)'),
    ('의료법 제56조', '의료광고 금지 — 치료 효과 보장, 의료기관이 아닌 자의 과장 광고 금지'),
    ('소비자보호법 제21조', '허위·과장 표시·광고 금지'),
    ('대법원 2010도1573', '\'진단\'이라는 명칭이 일반인에게 의료행위로 오인될 소지가 있으면 위법 소지'),
]
for law, desc in laws:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r_law = p.add_run(f'· [{law}]  ')
    r_law.font.bold = True
    r_law.font.size = Pt(10.5)
    r_law.font.name = '맑은 고딕'
    r_law.font.color.rgb = RGBColor(0,84,166)
    r_desc = p.add_run(desc)
    r_desc.font.size = Pt(10.5)
    r_desc.font.name = '맑은 고딕'

add_blank(doc)

# ══════════════════════════════════════════
# 2. 고위험 요소
# ══════════════════════════════════════════
add_heading(doc, '2. 고위험 요소 (즉시 수정 권장)', level=2)

# 2-1
add_heading(doc, '2-1. 앱 전반의 "진단" 용어 사용', level=3)
add_body(doc, '▶ 현재 사용 중인 위험 표현:', color=(180,0,0))
for txt in [
    '"AI 신체 밸런스 & 뇌 건강 진단 시스템"',
    '"BTC 3바디 AI 진단 센터"',
    '"통합 AI 리포트", "정밀 분석"',
]:
    add_body(doc, f'  · {txt}', indent=True)
add_blank(doc)
add_body(doc, '▶ 위험 이유:', color=(180,0,0))
add_body(doc, '"진단"은 의료법 제2조가 명시하는 의료행위 핵심 용어입니다. 무면허로 "진단"을 수행하는 것처럼 보이면 의료법 제27조 위반으로 3년 이하 징역 또는 3천만 원 이하 벌금이 부과될 수 있습니다.', indent=True)
add_blank(doc)
add_body(doc, '▶ 수정 방향:', color=(0,120,0))
for txt in [
    '"진단" → "측정", "평가", "체력 스크리닝", "웰니스 체크"로 전환',
    '예: "BTC 3바디 AI 체력 측정 센터", "AI 신체 균형 평가 시스템"',
]:
    add_body(doc, f'  · {txt}', indent=True)
add_blank(doc)

# 2-2
add_heading(doc, '2-2. "뇌 건강 진단" 표현 (가장 위험)', level=3)
add_body(doc, '▶ 현재 상태:', color=(180,0,0))
add_body(doc, 'BrainTestModule 컴포넌트 존재 + AI 프롬프트에 brainHealthImplication 항목 포함', indent=True)
add_blank(doc)
add_body(doc, '▶ 위험 이유:', color=(180,0,0))
add_body(doc, '뇌 기능 평가는 신경과·정신건강의학과의 전속 의료행위 영역입니다. "뇌 건강 진단"이라는 표현은 의료인 면허 없이 신경학적 진단을 수행하는 것으로 오인될 수 있어 고위험 표현에 해당합니다.', indent=True)
add_blank(doc)
add_body(doc, '▶ 수정 방향:', color=(0,120,0))
add_body(doc, '"뇌 건강 진단" → "인지 반응 측정", "집중력 체크", "두뇌 활성도 참고 지표"', indent=True)
add_blank(doc)

# 2-3
add_heading(doc, '2-3. AI 프롬프트 내 과장 효과 주장 (geminiService.ts 라인 366)', level=3)
add_body(doc, '▶ 현재 프롬프트 원문:', color=(180,0,0))
box = doc.add_paragraph()
box.paragraph_format.left_indent  = Cm(1)
box.paragraph_format.right_indent = Cm(1)
box.paragraph_format.space_before = Pt(4)
box.paragraph_format.space_after  = Pt(4)
r_box = box.add_run(
    '"신체의 불균형과 노화를 뇌 신경망 및 3바디 7코드의 막힘 현상으로 연결하고,\n'
    '\'광명차크라 수련\'이 모든 것의 근본 해결책임을 강력하게 어필하세요."'
)
r_box.font.size = Pt(10)
r_box.font.name = '맑은 고딕'
r_box.font.color.rgb = RGBColor(120,0,0)
r_box.font.italic = True
add_blank(doc)
add_body(doc, '▶ 위험 이유:', color=(180,0,0))
for txt in [
    '의료법 제56조: 질병·신체 상태와 특정 치료법의 인과관계를 단정하는 표현 금지',
    '소비자보호법: "근본 해결책"은 효능·효과를 보장하는 허위 과장 광고에 해당',
    '신체 검사 데이터를 근거로 특정 프로그램 구매를 유도하는 구조는 민법상 기망 행위 성립 가능',
]:
    add_body(doc, f'  · {txt}', indent=True)
add_blank(doc)
add_body(doc, '▶ 수정 방향:', color=(0,120,0))
for txt in [
    '"근본 해결책임을 강력하게 어필" → "개선에 도움이 될 수 있는 활동으로 안내"',
    '"막힘 현상" 등 의료적 단정 표현 삭제',
]:
    add_body(doc, f'  · {txt}', indent=True)
add_blank(doc)

# 2-4
add_heading(doc, '2-4. 비만 진단 및 처방성 표현 (geminiService.ts 라인 365)', level=3)
add_body(doc, '▶ 현재 프롬프트 원문:', color=(180,0,0))
add_body(doc, '"과체중" 또는 "비만"으로 파악되면 "체지방 감량, 심혈관 대사 관리, 관절 부하 솔루션" 포함 지시', indent=True)
add_blank(doc)
add_body(doc, '▶ 위험 이유:', color=(180,0,0))
add_body(doc, '비만 진단 및 심혈관·대사 질환 관리 처방은 의료행위입니다. AI가 체형 분석 결과를 토대로 처방성 솔루션을 제시하면 의료법 위반 소지가 있습니다.', indent=True)
add_blank(doc)
add_body(doc, '▶ 수정 방향:', color=(0,120,0))
for txt in [
    '"심혈관 대사 관리" 등 처방성 표현 삭제',
    '"전문 의료기관 상담을 권장합니다" 형태의 리퍼럴 표현으로 대체',
]:
    add_body(doc, f'  · {txt}', indent=True)
add_blank(doc)

# ══════════════════════════════════════════
# 3. 주의 요소 — 표
# ══════════════════════════════════════════
add_heading(doc, '3. 주의 요소 (표현 조정 권장)', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

headers = ['항목', '현재 표현', '위험도', '권장 수정 방향']
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255,255,255)
    # 헤더 배경색
    tc = hdr_cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '003366')
    tcPr.append(shd)

rows_data = [
    ('안면 분석', '피부 노화·주름 탄력 분석', '⚠️ 중간', '상담 보조 참고 지표 명시 (현재 프롬프트 일부 있음 ✅)'),
    ('자세 분석', '심한 거북목 — 노화 신호', '⚠️ 중간', '거북목 패턴 관찰 (단정형 → 관찰형으로 전환)'),
    ('점수 감점', '경각심 위해 가차없이 감점', '⚠️ 중간', '삭제 — 의도적 과소평가는 허위정보 제공'),
    ('리포트 언어', 'AI 생성 질환·위험 표현', '🟢 낮음', '현재 프롬프트 금지어 목록 운영 중 ✅'),
]

for row_data in rows_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        row_cells[i].paragraphs[0].runs[0].font.name = '맑은 고딕'
        row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

add_blank(doc)

# ══════════════════════════════════════════
# 4. 현재 준수 사항
# ══════════════════════════════════════════
add_heading(doc, '4. 현재 준수 사항 (유지)', level=2)
compliant = [
    'AI 프롬프트 내 면책 고지 존재: "본 시스템은 의료용 진단기기가 아닌 체력 및 건강 증진용 웰니스 스크리닝 도구입니다."',
    'AI 금지어 목록 운영: "스트레스가 높습니다", "불안정합니다", "질환" 등 생성 금지',
    '안면 분석 프롬프트에 "임상적 피부과 진단이 아닌 상담 보조용 참고 지표" 명시',
    '국민체력100 기반 체력 검사 항목(균형·유연성·스쿼트·푸시업)은 비의료 영역 — 위법 소지 낮음',
    '측정 데이터 기반 엄격한 코드 확정값 사용으로 AI 임의 과장 방지',
]
for item in compliant:
    add_body(doc, f'  ✅  {item}', indent=True, color=(0,100,0))
add_blank(doc)

# ══════════════════════════════════════════
# 5. 우선순위별 수정 계획
# ══════════════════════════════════════════
add_heading(doc, '5. 우선순위별 수정 계획', level=2)

add_heading(doc, '[즉시 조치] 법적 위험 항목', level=3)
immediate = [
    '앱 내 UI 전체: "진단" → "측정/평가/스크리닝"으로 일괄 교체',
    'AI 프롬프트: "광명차크라가 모든 것의 근본 해결책" 문구 삭제 또는 완화',
    '"뇌 건강 진단" → "인지 반응 측정"으로 명칭 변경',
    '비만 처방성 표현(심혈관 대사 관리 등) → 전문의 상담 권장 문구로 대체',
]
for i, item in enumerate(immediate, 1):
    add_body(doc, f'  {i}.  {item}', indent=True, color=(180,0,0))
add_blank(doc)

add_heading(doc, '[권장 조치] 신뢰도 향상 항목', level=3)
recommended = [
    '리포트 화면 하단에 법적 면책 문구 UI 표시 (현재 프롬프트에만 있고 화면에 미노출)',
    '프로그램 추천 섹션에 "본 추천은 체력 측정 결과 참고 제안이며 의료적 처방이 아닙니다" 추가',
    '"정밀 분석" → "참고 분석"으로 표현 완화',
    '최종 배포 전 의료법 전문 변호사 검토 수검',
]
for i, item in enumerate(recommended, 5):
    add_body(doc, f'  {i}.  {item}', indent=True, color=(0,84,166))
add_blank(doc)

# ══════════════════════════════════════════
# 6. 면책 문구 예시
# ══════════════════════════════════════════
add_heading(doc, '6. 법적 면책 문구 권장 예시 (화면 표시용)', level=2)

disclaimer_p = doc.add_paragraph()
disclaimer_p.paragraph_format.left_indent  = Cm(1)
disclaimer_p.paragraph_format.right_indent = Cm(1)
disclaimer_p.paragraph_format.space_before = Pt(6)
disclaimer_p.paragraph_format.space_after  = Pt(6)

r_disc = disclaimer_p.add_run(
    '⚠️  본 서비스 이용 안내\n\n'
    '본 측정 결과는 체력 및 건강 증진 목적의 참고 지표이며,\n'
    '의료기기법에 따른 의료기기 또는 의료행위가 아닙니다.\n'
    '건강 이상이 의심될 경우 반드시 의료기관을 방문하시기 바랍니다.'
)
r_disc.font.name = '맑은 고딕'
r_disc.font.size = Pt(10)
r_disc.font.color.rgb = RGBColor(100,50,0)

# 배경 박스 테두리
pPr = disclaimer_p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','left','bottom','right']:
    bdr = OxmlElement(f'w:{side}')
    bdr.set(qn('w:val'), 'single')
    bdr.set(qn('w:sz'), '6')
    bdr.set(qn('w:space'), '4')
    bdr.set(qn('w:color'), 'CC8800')
    pBdr.append(bdr)
pPr.append(pBdr)

add_blank(doc)

# ══════════════════════════════════════════
# 7. 결론
# ══════════════════════════════════════════
add_heading(doc, '7. 결론 및 권고', level=2)
add_body(doc, 'BTC 3바디 AI분석기 v1.8.4는 체력 측정 기능 자체는 국민체력100 등 공인 비의료 서비스와 유사한 구조를 가지고 있어 기능 측면의 위법 소지는 낮습니다.')
add_blank(doc)
add_body(doc, '그러나 다음 3가지 요소가 의료법·소비자보호법 위반 위험을 높이고 있습니다:', color=(180,0,0))
for txt in [
    '"진단"이라는 용어의 앱 전반 광범위 사용',
    '특정 프로그램을 "근본 해결책"으로 단정하는 AI 생성 문구',
    '"뇌 건강 진단"이라는 의료 영역 직접 침범 표현',
]:
    add_body(doc, f'  · {txt}', indent=True, color=(180,0,0))
add_blank(doc)
add_body(doc, '위 3가지를 상용 배포 전 수정하면 법적 리스크를 크게 낮출 수 있습니다.', color=(0,100,0))
add_body(doc, '최종 배포 전 의료법 전문 변호사의 최종 검토를 받을 것을 권고합니다.', color=(0,84,166))

# ── 저장 ──
output_path = r'd:\antigravity_vibecoding\BT 3바디 ai테스트\docs\BTC_AI분석기_의료법_검토보고서.docx'
doc.save(output_path)
print(f'완료: {output_path}')
