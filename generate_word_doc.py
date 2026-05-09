import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Add a Title
title = doc.add_heading('3바디 AI 건강나이 측정 서비스', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('법적 및 학술적 산출 근거 (마스터 가이드)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1
doc.add_heading('1. 서비스 정의 및 법적 위치 (의료법 위반 방지 가이드)', level=1)
p1 = doc.add_paragraph('3바디 AI 테스트는 의료기관의 진단이나 처방을 목적으로 하는 의료기기가 아닙니다. 보건복지부의 ')
p1.add_run('「비의료 건강관리서비스 가이드라인」').bold = True
p1.add_run('을 엄격히 준수하며, 일상적인 건강 유지와 질병의 사전 예방을 돕기 위한 ')
p1.add_run('\'웰니스(Wellness) 스크리닝 및 동기부여\'').bold = True
p1.add_run(' 서비스입니다.')

bullet1 = doc.add_paragraph(style='List Bullet')
bullet1.add_run('마케팅 홍보 시 절대 금칙어: ').bold = True
bullet1.add_run('진단, 진찰, 치료, 처방, 환자, 특정 병명(거북목증후군, 치매 등) 예측 및 완치 보장')

bullet2 = doc.add_paragraph(style='List Bullet')
bullet2.add_run('권장 대체 표현: ').bold = True
bullet2.add_run('분석, 추정, 측정 테스트, 웰니스 스크리닝, 건강 가이드, 맞춤 솔루션, 체형 불균형 확인, 인지 반응 체크')

doc.add_paragraph()

# Section 2
doc.add_heading('2. 연령 지표(건강나이) 종합 산출 원리', level=1)
p2 = doc.add_paragraph('본 서비스에서 제공하는 건강나이(신체, 뇌, 얼굴, 마음)는 임의로 만들어진 수치가 아닙니다. Vision AI를 통한 측정 데이터(관절 각도, 인지 반응 속도, 안면 텍스처 등)를 ')
p2.add_run('세계적으로 검증된 인지과학 및 생체역학의 학술적 연구 방법론').bold = True
p2.add_run('과 결합하여 당사 고유의 통계적·휴리스틱 알고리즘으로 분석한 결과입니다.')

doc.add_paragraph()

# Section 3
doc.add_heading('3. 세부 항목별 학술적 측정 기반 (Reference)', level=1)

# 3.1
doc.add_heading('신체나이 (Physical Age) : 구조적 정렬 및 자세 균형', level=2)
p3_1 = doc.add_paragraph(style='List Bullet')
p3_1.add_run('측정 원리: ').bold = True
p3_1.add_run('카메라 비전 AI가 신체의 관절 랜드마크를 추출하여, 목-어깨-골반의 비대칭성과 수직 중심축(Plumb Line)의 이탈 정도를 역학적으로 분석합니다.')
p3_1_2 = doc.add_paragraph(style='List Bullet')
p3_1_2.add_run('학술적 근거: ').bold = True
p3_1_2.add_run('Kendall, F. P. (2005)의 자세 평가(Posture Analysis) 모델\n')
p3_1_2.add_run('Reference: "Muscles: Testing and function with posture and pain." 전 세계 물리치료 및 근골격계 불균형 평가의 글로벌 표준 방법론을 채택했습니다.').italic = True

# 3.2
doc.add_heading('뇌나이 (Brain Age) : 두뇌 인지 반응 및 작업기억', level=2)
p3_2 = doc.add_paragraph(style='List Bullet')
p3_2.add_run('측정 원리: ').bold = True
p3_2.add_run('단순 지능이 아닌, 전두엽의 실행 제어 능력(충동 억제)과 단기 기억력을 밀리초(ms) 단위의 반응 속도와 정답률로 측정합니다.')
p3_2_2 = doc.add_paragraph(style='List Bullet')
p3_2_2.add_run('학술적 근거 1: ').bold = True
p3_2_2.add_run('J. R. Stroop (1935)의 인지 간섭 현상 연구\n')
p3_2_2.add_run('Reference: "Studies of interference in serial verbal reactions." 글자와 색상의 불일치 상황에서 뇌가 겪는 부하(Stroop Effect)를 통해 전두엽의 반응 및 억제력을 평가합니다.').italic = True
p3_2_3 = doc.add_paragraph(style='List Bullet')
p3_2_3.add_run('학술적 근거 2: ').bold = True
p3_2_3.add_run('A. Baddeley (1974)의 작업기억 모델\n')
p3_2_3.add_run('Reference: "Working memory." 뇌가 시공간적 정보를 일시적으로 저장하고 처리하는 능력을 평가하는 마트 장보기(기억+산술) 테스트의 기초가 되었습니다.').italic = True

# 3.3
doc.add_heading('얼굴나이 (Face Age) : 안면 랜드마크 텍스처 분석', level=2)
p3_3 = doc.add_paragraph(style='List Bullet')
p3_3.add_run('측정 원리: ').bold = True
p3_3.add_run('노화와 연관된 주요 안면 특징(피부 탄력 저하에 따른 윤곽선 하강, 주름의 음영 깊이, 좌우 비대칭성)을 딥러닝 기반으로 추출합니다.')
p3_3_2 = doc.add_paragraph(style='List Bullet')
p3_3_2.add_run('학술적 근거: ').bold = True
p3_3_2.add_run('Computer Vision 기반 안면 연령 추정(Facial Age Estimation) 모델\n')
p3_3_2.add_run('최신 AI 안면 인식 랜드마크 추출 기술을 사용하여, 동 연령대 대규모 표본 데이터의 시각적 노화 패턴과 사용자의 현재 안면 상태를 통계적으로 대조 추정합니다.').italic = True

# 3.4
doc.add_heading('마음나이 (Mind Age) : 정서 및 에너지 밸런스', level=2)
p3_4 = doc.add_paragraph(style='List Bullet')
p3_4.add_run('측정 원리: ').bold = True
p3_4.add_run('스트레스에 대한 신체적·심리적 저항력과 교감/부교감 신경의 활성도(에너지 흐름) 수준을 설문 및 측정 데이터를 통해 점수화합니다.')
p3_4_2 = doc.add_paragraph(style='List Bullet')
p3_4_2.add_run('학술적 근거: ').bold = True
p3_4_2.add_run('만성 스트레스 척도(PSS) 및 자율신경계(HRV) 모델링 기반\n')
p3_4_2.add_run('심리적 피로도와 회복탄력성을 측정하는 표준화된 심리학적 척도를 웰니스 관점의 에너지 연령 지표로 치환하여 적용했습니다.').italic = True

doc.add_paragraph()

# Section 4
doc.add_heading('4. 실제 어플리케이션 내 적용된 법적 면책 조항', level=1)
p4 = doc.add_paragraph('(현재 온라인 웹 리포트 하단에 실시간으로 반영되어 출력되는 문구로, 법적 방어막 역할을 합니다.)\n')

quote = doc.add_paragraph(style='Quote')
r = quote.add_run('[서비스 산출 근거 및 이용 안내 (법적 고지)]\n\n')
r.bold = True
quote.add_run('📌 학술적 측정 기반: 본 서비스의 웰니스 지표는 Kendall의 자세 평가(Posture Analysis), J.R. Stroop(1935)의 인지 간섭 현상, A. Baddeley(1974)의 작업기억 모델, Computer Vision 기반 안면 랜드마크 기술 등 검증된 인지과학 및 생체역학 연구 방법론을 기초로 AI 알고리즘화 되었습니다.\n\n')
quote.add_run('📌 연령 지표 산출 근거: 제공되는 신체/뇌/얼굴/마음 건강나이는 질병 진단을 위한 수치가 아니며, 측정 데이터를 종합하여 통계적·휴리스틱 알고리즘으로 산출한 건강관리 목적의 참고 지표입니다.\n\n')
quote.add_run('⚠️ 비의료 건강관리서비스 안내: 본 테스트는 보건복지부의 가이드라인을 준수합니다. 본 결과는 의료법에 따른 의료행위(진찰, 검사, 처방 등)를 대체할 수 없으며, 질환이 의심될 경우 반드시 전문 의료기관을 방문하시기 바랍니다.')

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.add_run('이 마스터 가이드를 바탕으로 "세계적 연구 논문에 기반한 검증된 웰니스 AI를 무료로 만나보세요!"라는 강력하고 안전한 마케팅 메시지를 구성하실 수 있습니다.').bold = True

doc.save('3바디_AI_건강나이_마스터가이드.docx')
print('Word document saved successfully.')
